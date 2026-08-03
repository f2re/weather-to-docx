from __future__ import annotations

from collections.abc import Iterable
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from weather_to_docx.document.icons import WeatherIconRenderer
from weather_to_docx.document.styles import (
    DANGER,
    DARK_BLUE,
    LIGHT_BLUE,
    LIGHT_GREY,
    MID_BLUE,
    WARNING,
    WHITE,
    configure_document,
    prevent_row_split,
    set_cell_shading,
    set_cell_text,
    set_cell_width,
    set_repeat_header_count,
    set_table_fixed_layout,
)
from weather_to_docx.document.weather_rules import weather_presentation
from weather_to_docx.domain.models import (
    DocumentOptions,
    ForecastPoint,
    ForecastSeries,
    Location,
    QualityFlag,
)
from weather_to_docx.domain.parameters import definition
from weather_to_docx.utils.meteorology import wind_rumb

QUALITY_MARKERS = {
    QualityFlag.INTERPOLATED: "≈",
    QualityFlag.CALCULATED: "*",
    QualityFlag.CORRECTED: "†",
    QualityFlag.STALE: "!",
    QualityFlag.SUSPECT: "?",
    QualityFlag.MISSING: "—",
}


class DocumentGenerator:
    def __init__(self, icon_cache_dir: Path) -> None:
        self.icons = WeatherIconRenderer(icon_cache_dir)

    def generate(
        self,
        *,
        location: Location,
        series: list[ForecastSeries],
        options: DocumentOptions,
        output_path: Path,
    ) -> Path:
        if not series:
            raise ValueError("Для документа не передан ни один прогностический ряд")
        mismatched = [item.location.id for item in series if item.location.id != location.id]
        if mismatched:
            raise ValueError("В документ нельзя объединять прогнозы для разных координат")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        configure_document(document, options.page_size)
        self._configure_properties(document, location, options)
        self._add_header(document, location, series, options)

        for index, forecast in enumerate(series):
            if index:
                document.add_page_break()
            self._add_source_section(document, forecast, options)

        self._add_footer(document, location)
        document.save(output_path)
        return output_path

    @staticmethod
    def _configure_properties(document: Document, location: Location, options: DocumentOptions) -> None:
        properties = document.core_properties
        properties.title = f"{options.title}: {location.name}"
        properties.subject = "Автоматически сформированный метеорологический прогноз"
        properties.author = options.prepared_by or "weather-to-docx"
        properties.keywords = "погода, прогноз, GFS, DOCX, метеорология"
        properties.comments = "Сформировано системой weather-to-docx без макросов и внешних ресурсов."

    def _add_header(
        self,
        document: Document,
        location: Location,
        series: list[ForecastSeries],
        options: DocumentOptions,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(options.title)
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(18)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(location.name)
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(13)

        if options.organisation:
            organisation = document.add_paragraph(options.organisation)
            organisation.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.add_table(rows=3, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        set_table_fixed_layout(table)
        values = [
            ("Координаты", f"{location.latitude:.5f}, {location.longitude:.5f}"),
            ("Высота", f"{location.elevation_m:.0f} м" if location.elevation_m is not None else "не задана"),
            ("Часовой пояс", location.timezone),
            ("Сформировано", datetime.now(UTC).strftime("%d.%m.%Y %H:%M UTC")),
            ("Источники", ", ".join(forecast.source.source_id for forecast in series)),
            ("Количество рядов", str(len(series))),
        ]
        for index, (label, value) in enumerate(values):
            row = index // 2
            column = (index % 2) * 2
            set_cell_text(table.cell(row, column), label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(table.cell(row, column), MID_BLUE)
            set_cell_text(table.cell(row, column + 1), value, align=WD_ALIGN_PARAGRAPH.LEFT)

        warning = document.add_paragraph()
        warning.paragraph_format.space_before = Pt(5)
        warning.paragraph_format.space_after = Pt(7)
        warning_run = warning.add_run(
            "Прогноз является расчётной информацией. Перед принятием критически важных решений "
            "необходимо учитывать наблюдения, предупреждения уполномоченной метеослужбы и неопределённость модели."
        )
        warning_run.bold = True
        warning_run.font.size = Pt(8.5)

    def _add_source_section(
        self,
        document: Document,
        forecast: ForecastSeries,
        options: DocumentOptions,
    ) -> None:
        source = forecast.source
        heading = document.add_heading(level=1)
        heading.add_run(f"{source.provider} — {source.model}")

        metadata = document.add_table(rows=3, cols=4)
        metadata.style = "Table Grid"
        metadata.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_fixed_layout(metadata)
        cycle = (
            source.cycle_time_utc.astimezone(UTC).strftime("%d.%m.%Y %H:%M UTC")
            if source.cycle_time_utc
            else "не указан поставщиком"
        )
        pairs = [
            ("Источник", source.source_id),
            ("Цикл", cycle),
            ("Получено", source.retrieved_at_utc.astimezone(UTC).strftime("%d.%m.%Y %H:%M UTC")),
            ("Горизонт", f"{source.horizon_hours} ч" if source.horizon_hours is not None else "—"),
            ("Шаг", f"{source.native_time_step_hours:g} ч" if source.native_time_step_hours else "переменный"),
            ("Сетка", " / ".join(filter(None, (source.spatial_resolution, source.grid_type))) or "—"),
        ]
        for index, (label, value) in enumerate(pairs):
            row = index // 2
            column = (index % 2) * 2
            set_cell_text(metadata.cell(row, column), label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(metadata.cell(row, column), MID_BLUE)
            set_cell_text(metadata.cell(row, column + 1), value, align=WD_ALIGN_PARAGRAPH.LEFT)

        document.add_heading("1. Наглядный прогноз", level=2)
        self._add_summary_table(document, forecast, options)

        if options.include_detailed_table:
            document.add_heading("2. Подробный почасовой метеорологический отчёт", level=2)
            self._add_detailed_table(document, forecast)

        notes = list(forecast.warnings)
        if source.grid_distance_km is not None:
            notes.append(f"Расстояние до ближайшего модельного узла: {source.grid_distance_km:.1f} км.")
        if source.attribution:
            notes.append(f"Атрибуция: {source.attribution}.")
        if source.licence:
            notes.append(f"Условия использования: {source.licence}.")
        if source.source_reference:
            notes.append(f"Технический источник: {source.source_reference}.")
        notes.append("Обозначения качества: ≈ интерполяция; * расчётное значение; † исправлено контролем качества; ! устаревшее; ? сомнительное.")

        for text in notes:
            paragraph = document.add_paragraph(style=None)
            paragraph.paragraph_format.left_indent = Mm(3)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(f"• {text}")
            run.font.size = Pt(7.5)

    def _add_summary_table(
        self,
        document: Document,
        forecast: ForecastSeries,
        options: DocumentOptions,
    ) -> None:
        points = list(self._summary_points(forecast.points, options))
        headers = (
            "Дата и время",
            "Погода",
            "Температура",
            "Ощущается",
            "Осадки",
            "Вероятность",
            "Ветер",
            "Порывы",
            "Давление",
            "Облачность",
        )
        widths = (30, 34, 20, 20, 24, 22, 36, 22, 24, 24)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(table)
        for cell, text, width in zip(table.rows[0].cells, headers, widths, strict=True):
            set_cell_text(cell, text, size=8, bold=True)
            set_cell_shading(cell, DARK_BLUE)
            self._set_text_color(cell, WHITE)
            set_cell_width(cell, width)
        set_repeat_header_count(table, 1)

        previous_date = None
        for point in points:
            row = table.add_row()
            prevent_row_split(row)
            presentation = weather_presentation(point)
            current_date = point.valid_time_local.date()
            row_fill = LIGHT_BLUE if current_date != previous_date else WHITE
            previous_date = current_date
            for cell, width in zip(row.cells, widths, strict=True):
                set_cell_width(cell, width)
                set_cell_shading(cell, row_fill)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            set_cell_text(row.cells[0], point.valid_time_local.strftime("%d.%m.%Y\n%H:%M"), size=7.5)
            self._set_icon_cell(row.cells[1], presentation.icon_key, presentation.description)
            set_cell_text(row.cells[2], self._value(point, "temperature_2m", with_unit=True), size=8, bold=True)
            set_cell_text(row.cells[3], self._value(point, "apparent_temperature", with_unit=True), size=8)
            set_cell_text(row.cells[4], self._value(point, "precipitation", with_unit=True), size=8)
            set_cell_text(row.cells[5], self._value(point, "precipitation_probability", with_unit=True), size=8)
            set_cell_text(row.cells[6], self._wind_text(point), size=7.5)
            set_cell_text(row.cells[7], self._value(point, "wind_gusts_10m", with_unit=True), size=8)
            set_cell_text(row.cells[8], self._value(point, "pressure_msl", with_unit=True), size=8)
            set_cell_text(row.cells[9], self._value(point, "cloud_cover", with_unit=True), size=8)
            self._apply_hazard_shading(row.cells, point)

    def _add_detailed_table(self, document: Document, forecast: ForecastSeries) -> None:
        headers = (
            "Дата и время",
            "Срок",
            "Погода",
            "T, °C",
            "Ощ., °C",
            "Td, °C",
            "RH, %",
            "Давление, гПа",
            "Ветер 10 м",
            "Порывы",
            "Осадки, мм",
            "PoP, %",
            "Снег",
            "Облачность, %",
            "Видимость",
            "Конвекция",
            "Радиация",
            "T почвы, °C",
            "Влажность почвы",
            "ET₀ / ПС",
        )
        widths = (28, 12, 25, 11, 11, 11, 11, 25, 30, 14, 30, 12, 23, 32, 18, 25, 31, 28, 34, 22)
        groups = (
            ("Время", 0, 1),
            ("Явление", 2, 2),
            ("Температура и влажность", 3, 6),
            ("Давление", 7, 7),
            ("Ветер", 8, 9),
            ("Осадки", 10, 12),
            ("Облачность", 13, 14),
            ("Конвекция", 15, 15),
            ("Радиация", 16, 16),
            ("Почва и ПС", 17, 19),
        )
        table = document.add_table(rows=2, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(table)

        for group_name, start, end in groups:
            cell = table.cell(0, start)
            if end > start:
                cell = cell.merge(table.cell(0, end))
            set_cell_text(cell, group_name, size=7, bold=True)
            set_cell_shading(cell, DARK_BLUE)
            self._set_text_color(cell, WHITE)
        for cell, text, width in zip(table.rows[1].cells, headers, widths, strict=True):
            set_cell_text(cell, text, size=6.5, bold=True)
            set_cell_shading(cell, MID_BLUE)
            set_cell_width(cell, width)
        set_repeat_header_count(table, 2)

        previous_date = None
        for point in forecast.points:
            row = table.add_row()
            prevent_row_split(row)
            current_date = point.valid_time_local.date()
            fill = LIGHT_BLUE if current_date != previous_date else WHITE
            previous_date = current_date
            for cell, width in zip(row.cells, widths, strict=True):
                set_cell_width(cell, width)
                set_cell_shading(cell, fill)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            presentation = weather_presentation(point)
            cells = (
                point.valid_time_local.strftime("%d.%m.%Y\n%H:%M"),
                f"+{point.lead_hours} ч" if point.lead_hours is not None else "—",
                presentation.description,
                self._value(point, "temperature_2m"),
                self._value(point, "apparent_temperature"),
                self._value(point, "dew_point_2m"),
                self._value(point, "relative_humidity_2m"),
                self._pressure_text(point),
                self._wind_text(point),
                self._value(point, "wind_gusts_10m", with_unit=True),
                self._precipitation_text(point),
                self._value(point, "precipitation_probability"),
                self._snow_text(point),
                self._cloud_text(point),
                self._visibility_text(point),
                self._convection_text(point),
                self._radiation_text(point),
                self._soil_temperature_text(point),
                self._soil_moisture_text(point),
                self._surface_exchange_text(point),
            )
            for cell, text in zip(row.cells, cells, strict=True):
                set_cell_text(cell, text, size=6.2)
            self._apply_hazard_shading(row.cells, point)

    @staticmethod
    def _summary_points(points: list[ForecastPoint], options: DocumentOptions) -> Iterable[ForecastPoint]:
        if not points:
            return []
        selected: list[ForecastPoint] = []
        for index, point in enumerate(points):
            lead = point.lead_hours
            interval = (
                options.summary_interval_hours
                if lead is None or lead <= options.summary_switch_hour
                else options.extended_summary_interval_hours
            )
            if index == 0 or index == len(points) - 1 or lead is None or lead % interval == 0:
                selected.append(point)
        return selected

    def _set_icon_cell(self, cell, icon_key: str, description: str) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        shape = run.add_picture(str(self.icons.render(icon_key)), width=Mm(8))
        # Accessible description is embedded in wp:docPr and does not depend on
        # the icon filename or an external resource.
        shape._inline.docPr.set("descr", f"Пиктограмма погоды: {description}")
        shape._inline.docPr.set("title", description)
        description_run = paragraph.add_run(f"\n{description}")
        description_run.font.name = "Liberation Sans"
        description_run.font.size = Pt(6.5)

    @staticmethod
    def _set_text_color(cell, color: str) -> None:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                color_element = OxmlElement("w:color")
                color_element.set(qn("w:val"), color)
                run._r.get_or_add_rPr().append(color_element)

    @staticmethod
    def _marker(point: ForecastPoint, code: str) -> str:
        measurement = point.measurement(code)
        return QUALITY_MARKERS.get(measurement.quality, "") if measurement else ""

    def _value(self, point: ForecastPoint, code: str, *, with_unit: bool = False) -> str:
        measurement = point.measurement(code)
        if measurement is None or measurement.value is None:
            return "—"
        parameter = definition(code)
        value = measurement.value
        if isinstance(value, bool):
            text = "да" if value else "нет"
        elif isinstance(value, (int, float)):
            text = f"{float(value):.{parameter.precision}f}"
        else:
            text = str(value)
        marker = QUALITY_MARKERS.get(measurement.quality, "")
        unit = measurement.unit or parameter.default_unit
        return f"{text}{marker} {unit}".strip() if with_unit else f"{text}{marker}"

    def _pressure_text(self, point: ForecastPoint) -> str:
        return f"MSL {self._value(point, 'pressure_msl')}\nSFC {self._value(point, 'surface_pressure')}"

    def _wind_text(self, point: ForecastPoint) -> str:
        speed = _number(point.raw("wind_speed_10m"))
        direction = _number(point.raw("wind_direction_10m"))
        if speed is None:
            return "—"
        rumb = wind_rumb(direction, speed)
        direction_text = f"{direction:.0f}°" if direction is not None else "—"
        marker = self._marker(point, "wind_speed_10m") or self._marker(point, "wind_direction_10m")
        return f"{speed:.1f}{marker} м/с\nиз {rumb}, {direction_text}"

    def _precipitation_text(self, point: ForecastPoint) -> str:
        return (
            f"Σ {self._value(point, 'precipitation')}\n"
            f"дождь {self._value(point, 'rain')}\n"
            f"ливни {self._value(point, 'showers')}\n"
            f"конв. {self._value(point, 'convective_precipitation')}"
        )

    def _snow_text(self, point: ForecastPoint) -> str:
        return f"снег {self._value(point, 'snowfall')} см\nпокров {self._value(point, 'snow_depth')} м"

    def _cloud_text(self, point: ForecastPoint) -> str:
        return (
            f"общ {self._value(point, 'cloud_cover')}\n"
            f"Н/С/В {self._value(point, 'cloud_cover_low')}/"
            f"{self._value(point, 'cloud_cover_mid')}/"
            f"{self._value(point, 'cloud_cover_high')}"
        )

    def _visibility_text(self, point: ForecastPoint) -> str:
        value = _number(point.raw("visibility"))
        if value is None:
            return "—"
        marker = self._marker(point, "visibility")
        return f"{value / 1000:.1f}{marker} км" if value >= 1000 else f"{value:.0f}{marker} м"

    def _convection_text(self, point: ForecastPoint) -> str:
        return (
            f"CAPE {self._value(point, 'cape')}\n"
            f"CIN {self._value(point, 'cin')}\n"
            f"VPD {self._value(point, 'vapour_pressure_deficit')}"
        )

    def _radiation_text(self, point: ForecastPoint) -> str:
        sunshine = _number(point.raw("sunshine_duration"))
        sunshine_text = "—" if sunshine is None else f"{sunshine / 60:.0f} мин"
        return (
            f"КВ {self._value(point, 'shortwave_radiation')}\n"
            f"пр/рас {self._value(point, 'direct_radiation')}/"
            f"{self._value(point, 'diffuse_radiation')}\n"
            f"солн {sunshine_text}"
        )

    def _soil_temperature_text(self, point: ForecastPoint) -> str:
        return (
            f"0/6 {self._value(point, 'soil_temperature_0cm')}/"
            f"{self._value(point, 'soil_temperature_6cm')}\n"
            f"18/54 {self._value(point, 'soil_temperature_18cm')}/"
            f"{self._value(point, 'soil_temperature_54cm')}"
        )

    def _soil_moisture_text(self, point: ForecastPoint) -> str:
        return (
            f"0–1 {self._value(point, 'soil_moisture_0_to_1cm')}\n"
            f"1–3 {self._value(point, 'soil_moisture_1_to_3cm')}\n"
            f"3–9 {self._value(point, 'soil_moisture_3_to_9cm')}\n"
            f"9–27 {self._value(point, 'soil_moisture_9_to_27cm')}\n"
            f"27–81 {self._value(point, 'soil_moisture_27_to_81cm')}"
        )

    def _surface_exchange_text(self, point: ForecastPoint) -> str:
        return (
            f"ET₀ {self._value(point, 'et0_fao_evapotranspiration')}\n"
            f"ET {self._value(point, 'evapotranspiration')}\n"
            f"ПС {self._value(point, 'boundary_layer_height')} м"
        )

    @staticmethod
    def _apply_hazard_shading(cells, point: ForecastPoint) -> None:
        precipitation = _number(point.raw("precipitation")) or 0
        gust = _number(point.raw("wind_gusts_10m")) or 0
        cape = _number(point.raw("cape")) or 0
        temperature = _number(point.raw("temperature_2m"))
        if precipitation >= 10 or gust >= 20 or cape >= 2000:
            for cell in cells:
                set_cell_shading(cell, DANGER)
        elif precipitation >= 3 or gust >= 14 or cape >= 1000 or (temperature is not None and (temperature <= -25 or temperature >= 35)):
            for cell in cells:
                set_cell_shading(cell, WARNING)

    @staticmethod
    def _add_footer(document: Document, location: Location) -> None:
        for section in document.sections:
            footer = section.footer
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(
                f"weather-to-docx · {location.name} · документ сформирован автоматически · "
            )
            run.font.name = "Liberation Sans"
            run.font.size = Pt(7)
            field_begin = OxmlElement("w:fldChar")
            field_begin.set(qn("w:fldCharType"), "begin")
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = " PAGE "
            field_end = OxmlElement("w:fldChar")
            field_end.set(qn("w:fldCharType"), "end")
            run._r.append(field_begin)
            run._r.append(instruction)
            run._r.append(field_end)


def _number(value) -> float | None:
    try:
        return float(value) if value is not None else None
    except (TypeError, ValueError):
        return None
