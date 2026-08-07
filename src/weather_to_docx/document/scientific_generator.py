"""Точка импорта профессионального генератора DOCX."""

import math
from datetime import datetime
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Mm, Pt
from matplotlib.axes import Axes
from matplotlib.figure import Figure

from weather_to_docx.document import audit_generator as _audit_generator
from weather_to_docx.document import compact_generator as _compact_generator
from weather_to_docx.document.consistent_controls import (
    consistent_control_point,
    shade_daily_hazard,
)
from weather_to_docx.document.consistent_summary import (
    build_consistent_risk_signals,
    build_highlights,
    consistent_daily_model_metrics,
    consistent_daily_presentation_point,
    daily_agreement_from_metrics,
)
from weather_to_docx.document.impact_labels import (
    daily_pressure_text,
    daily_temperature_text,
)
from weather_to_docx.document.plain_language import (
    daily_precipitation_metrics_text,
    daily_precipitation_text,
    daily_wind_text,
    detail_precipitation_text,
    detail_wind_text,
    ensemble_members_text,
    ensemble_precipitation_text,
    risk_support_text,
    risk_value_text,
)
from weather_to_docx.document.russian_labels import apply_russian_display_names
from weather_to_docx.document.source_names import source_display_name
from weather_to_docx.document.styles import (
    DARK_BLUE,
    LIGHT_BLUE,
    WHITE,
    prevent_row_split,
    set_cell_shading,
    set_cell_text,
    set_cell_width,
    set_repeat_header_count,
    set_table_fixed_layout,
)
from weather_to_docx.domain.models import ForecastSeries
from weather_to_docx.plotting.meteogram import _values
from weather_to_docx.plotting.professional_meteogram import WIND_ARROWS
from weather_to_docx.plotting.semantic_meteogram import SemanticMeteogramRenderer


class DocumentMeteogramRenderer(SemanticMeteogramRenderer):
    """Рендерер для DOCX с самодостаточными и понятными подписями."""

    def _new_professional_figure(self, title: str):
        figure, axes = super()._new_professional_figure(title)
        figure.set_size_inches(11.4, 5.55, forward=True)
        figure.subplots_adjust(hspace=0.38)
        return figure, axes

    def _finish_professional(
        self,
        figure: Figure,
        axes: tuple[Axes, ...],
        times: list[datetime],
        *,
        ensemble: bool,
    ) -> None:
        super()._finish_professional(figure, axes, times, ensemble=ensemble)
        if ensemble:
            for item in figure.texts:
                text = item.get_text()
                if "процентили" in text:
                    item.set_text(
                        text.replace(
                            "тёмная полоса — 25–75-й процентили; "
                            "светлая — 10–90-й",
                            "тёмная полоса — центральные 50 % вариантов; "
                            "светлая — центральные 80 % вариантов",
                        )
                    )
        figure.subplots_adjust(
            left=0.065,
            right=0.93,
            top=0.93,
            bottom=0.12,
            hspace=0.38,
        )

    def render_deterministic(
        self,
        forecast: ForecastSeries,
        output_path: Path,
        *,
        title: str | None = None,
    ) -> Path:
        return super().render_deterministic(
            forecast,
            output_path,
            title=self._complete_title(forecast, title),
        )

    def render_ensemble(
        self,
        forecast: ForecastSeries,
        output_path: Path,
        *,
        title: str | None = None,
    ) -> Path:
        return super().render_ensemble(
            forecast,
            output_path,
            title=self._complete_title(forecast, title),
        )

    @staticmethod
    def _complete_title(forecast: ForecastSeries, title: str | None) -> str:
        base = (title or source_display_name(forecast)).replace(
            "неопределённость прогноза",
            "разброс вариантов прогноза",
        )
        if not forecast.points:
            return f"{forecast.location.name} · {base}"
        start = forecast.points[0].valid_time_local
        end = forecast.points[-1].valid_time_local
        period = (
            f"{start:%d.%m.%Y}"
            if start.date() == end.date()
            else f"{start:%d.%m}–{end:%d.%m.%Y}"
        )
        return f"{forecast.location.name} · {base} · {period}"

    def _plot_temperature(
        self,
        axis: Axes,
        x: np.ndarray,
        forecast: ForecastSeries,
        *,
        ensemble: bool,
    ) -> None:
        super()._plot_temperature(axis, x, forecast, ensemble=ensemble)
        if not ensemble:
            return
        for line in axis.get_lines():
            if line.get_label() == "медиана":
                line.set_label("середина вариантов")
        self._legend_above(axis, ncol=3, fontsize=7.0)

    def _plot_precipitation(
        self,
        axis: Axes,
        x: np.ndarray,
        forecast: ForecastSeries,
        *,
        ensemble: bool,
    ) -> None:
        existing_axes = set(axis.figure.axes)
        super()._plot_precipitation(axis, x, forecast, ensemble=ensemble)
        axis.set_yticks((0.5, 2.0, 5.0, 10.0))
        if not ensemble:
            return
        probability_axes = [item for item in axis.figure.axes if item not in existing_axes]
        if not probability_axes:
            return
        probability_axis = probability_axes[-1]
        probability_axis.set_ylabel("варианты, %", fontsize=7.2, rotation=90, labelpad=8)
        for line in probability_axis.get_lines():
            label = line.get_label()
            if label.startswith("P ≥"):
                line.set_label(
                    label.replace("P ≥", "варианты с ≥") + "/интервал"
                )
        self._combined_legend_above(
            axis,
            probability_axis,
            ncol=5,
            fontsize=6.15,
            anchor_y=1.19,
        )

    def _plot_wind_pressure(
        self,
        axis: Axes,
        x: np.ndarray,
        forecast: ForecastSeries,
        *,
        ensemble: bool,
    ) -> None:
        existing_axes = set(axis.figure.axes)
        super()._plot_wind_pressure(axis, x, forecast, ensemble=ensemble)
        if not ensemble:
            return
        pressure_axes = [item for item in axis.figure.axes if item not in existing_axes]
        for line in axis.get_lines():
            if line.get_label() == "90-й процентиль порывов":
                line.set_label("порывы: 90 % вариантов не выше")
        if pressure_axes:
            self._combined_legend_above(
                axis,
                pressure_axes[-1],
                ncol=4,
                fontsize=6.45,
            )

    def _add_wind_direction_arrows(
        self,
        axis: Axes,
        x: np.ndarray,
        forecast: ForecastSeries,
    ) -> None:
        direction = _values(forecast, "wind_direction_10m")
        finite = np.flatnonzero(np.isfinite(direction))
        if finite.size == 0:
            return
        step = max(1, math.ceil(finite.size / 28))
        for index in finite[::step]:
            arrow_index = int(((direction[index] % 360) + 22.5) // 45) % 8
            axis.text(
                x[index],
                0.04,
                WIND_ARROWS[arrow_index],
                transform=axis.get_xaxis_transform(),
                ha="center",
                va="bottom",
                fontsize=8.5,
                fontfamily="DejaVu Sans",
                color="#2f5d46",
                fontweight="bold",
            )


def _add_meteogram_image(
    document: Document,
    image_path: Path,
    *,
    description: str,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = Mm(140)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Mm(276))
    shape._inline.docPr.set("descr", description)
    shape._inline.docPr.set("title", description)


def _translated_chart_note(document: Document, text: str) -> None:
    replacements = {
        (
            "Серые полосы обозначают ночь. Осадки показаны за исходный "
            "интервал; стрелки внизу показывают направление ветра."
        ): (
            "Серые полосы обозначают ночь. Столбики показывают интенсивность "
            "осадков, мм/ч; подписи над панелью — сумму за местные сутки. "
            "Стрелки внизу показывают направление ветра."
        ),
        (
            "Тёмная полоса — 25–75-й процентили, светлая — 10–90-й. "
            "Вероятности осадков относятся к указанным порогам и "
            "интервалам."
        ): (
            "Тёмная полоса показывает центральные 50 % вариантов ансамбля, "
            "светлая — центральные 80 %. Проценты справа у осадков — доля "
            "вариантов, в которых за один расчётный интервал выпало не меньше "
            "указанной суммы. Например, 5 % означает примерно 5 вариантов из "
            "100. Это не вероятность суточной суммы."
        ),
    }
    _ORIGINAL_ADD_CHART_NOTE(document, replacements.get(text, text))


def _replace_text_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    if paragraph.runs:
        paragraph.runs[0].text = paragraph.text.replace(old, new)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(paragraph.text.replace(old, new))
    return True


def _replace_visible_text(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        for old, new in replacements.items():
            _replace_text_in_paragraph(paragraph, old, new)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in replacements.items():
                        _replace_text_in_paragraph(paragraph, old, new)


def _plain_risk_cards(self, document: Document, risks) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run("Ключевые риски")
    run.bold = True
    run.font.name = "Liberation Sans"
    run.font.size = Pt(10.5)

    if not risks:
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        set_table_fixed_layout(table)
        set_cell_shading(table.cell(0, 0), _audit_generator.RISK_CLEAR)
        set_cell_text(
            table.cell(0, 0),
            "Существенных погодных рисков по выбранным моделям не выявлено.",
            size=8,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )
        return

    headers = ("Явление", "Когда", "Что ожидается", "По моделям")
    widths = (43, 66, 94, 64)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    for cell, text, width in zip(table.rows[0].cells, headers, widths, strict=True):
        set_cell_text(cell, text, size=7.2, bold=True)
        set_cell_shading(cell, DARK_BLUE)
        self._set_text_color(cell, WHITE)
        set_cell_width(cell, width)
    set_repeat_header_count(table, 1)

    for signal in risks:
        row = table.add_row()
        prevent_row_split(row)
        fill = (
            _audit_generator.RISK_HIGH
            if signal.severity >= 85
            else _audit_generator.RISK_MEDIUM
            if signal.severity >= 70
            else _audit_generator.RISK_LOW
        )
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        peak = (
            f"\nнаиболее выражено около {signal.peak_local:%H:%M}"
            if signal.peak_local is not None
            else ""
        )
        values = (
            signal.phenomenon,
            signal.time_text + peak,
            risk_value_text(signal.phenomenon, signal.value_text),
            risk_support_text(signal.support_count, signal.model_count),
        )
        for index, (cell, text) in enumerate(zip(row.cells, values, strict=True)):
            set_cell_text(
                cell,
                text,
                size=7.2,
                bold=index == 0,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )


def _plain_daily_table_professional(
    self,
    document: Document,
    forecasts: list[ForecastSeries],
    report_dates,
) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(3)
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run("Прогноз по дням")
    run.bold = True
    run.font.name = "Liberation Sans"
    run.font.size = Pt(10)

    headers = ("Дата", "Погода", "Температура", "Осадки", "Ветер", "Давление")
    widths = (24, 48, 36, 53, 59, 35)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    for cell, text, width in zip(table.rows[0].cells, headers, widths, strict=True):
        set_cell_text(cell, text, size=7.2, bold=True)
        set_cell_shading(cell, DARK_BLUE)
        self._set_text_color(cell, WHITE)
        set_cell_width(cell, width)
    set_repeat_header_count(table, 1)

    for day in report_dates:
        metrics = [
            metric
            for forecast in forecasts
            if (metric := consistent_daily_model_metrics(forecast, day)) is not None
        ]
        if not metrics:
            continue
        representative = consistent_daily_presentation_point(metrics, day)
        row = table.add_row()
        prevent_row_split(row)
        _compact_generator._set_row_height(row, 8)
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        set_cell_text(
            row.cells[0],
            f"{day:%d.%m}\n{_compact_generator._weekday_short(day)}",
            size=7.2,
            bold=True,
        )
        presentation = _audit_generator.weather_presentation(representative)
        self._set_compact_icon_cell(
            row.cells[1],
            presentation.icon_key,
            presentation.description,
        )
        set_cell_text(row.cells[2], daily_temperature_text(metrics), size=7)
        set_cell_text(row.cells[3], daily_precipitation_text(forecasts, day), size=7)
        set_cell_text(row.cells[4], daily_wind_text(metrics), size=7)
        set_cell_text(row.cells[5], daily_pressure_text(metrics), size=7)
        shade_daily_hazard(row.cells, metrics)


def _plain_compact_daily_table(self, document: Document, selection, report_dates) -> None:
    _plain_daily_table_professional(self, document, selection.usable, report_dates)


def _plain_ensemble_summary(self, document: Document, ensembles, report_dates) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(3)
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run("Ансамбль: разброс вариантов прогноза")
    run.bold = True
    run.font.name = "Liberation Sans"
    run.font.size = Pt(10)

    if not ensembles:
        paragraph = document.add_paragraph("Ансамбль не выбран.")
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            run.font.name = "Liberation Sans"
            run.font.size = Pt(7.2)
        return

    intro = document.add_paragraph(
        "Ансамбль — это много вариантов одного расчёта. Он показывает, "
        "насколько прогноз может меняться."
    )
    intro.paragraph_format.space_after = Pt(2)
    for run in intro.runs:
        run.font.name = "Liberation Sans"
        run.font.size = Pt(7.2)

    headers = (
        "Дата",
        "Ансамбль",
        "Температура\nу центральных 80 % вариантов",
        "Осадки\nдоля вариантов",
        "Порывы\n90 % вариантов не выше",
        "Рассчитано\nвариантов",
    )
    widths = (23, 34, 48, 72, 47, 34)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table)
    for cell, text, width in zip(table.rows[0].cells, headers, widths, strict=True):
        set_cell_text(cell, text, size=6.6, bold=True)
        set_cell_shading(cell, DARK_BLUE)
        self._set_text_color(cell, WHITE)
        set_cell_width(cell, width)
    _compact_generator._set_row_height(table.rows[0], 7)
    set_repeat_header_count(table, 1)

    row_count = 0
    for forecast in ensembles:
        for day in report_dates:
            points = [
                point
                for point in forecast.points
                if point.valid_time_local.date() == day
            ]
            if not points:
                continue
            row = table.add_row()
            prevent_row_split(row)
            _compact_generator._set_row_height(row, 7)
            for cell, width in zip(row.cells, widths, strict=True):
                set_cell_width(cell, width)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_shading(cell, LIGHT_BLUE if row_count % 2 == 0 else WHITE)
            row_count += 1
            noon = min(points, key=lambda point: abs(point.valid_time_local.hour - 12))
            values = (
                day.strftime("%d.%m"),
                source_display_name(forecast),
                _compact_generator._ensemble_temperature_text(noon),
                ensemble_precipitation_text(points),
                _compact_generator._ensemble_gust_text(points),
                ensemble_members_text(points, forecast),
            )
            for index, (cell, text) in enumerate(zip(row.cells, values, strict=True)):
                set_cell_text(cell, text, size=6.35, bold=index == 0)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(1)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "Например, 5 % вариантов означает: примерно 5 из 100 вариантов "
        "ансамбля дают не меньше указанной суммы осадков за один расчётный "
        "интервал. Это не вероятность суточной суммы. Температурный диапазон "
        "содержит центральные 80 % вариантов; для порывов показан уровень, "
        "который не превышают 90 % вариантов."
    )
    run.font.name = "Liberation Sans"
    run.font.size = Pt(6.6)


def _plain_compact_header(self, document, location, selection, ensembles, report_dates, options):
    _ORIGINAL_COMPACT_HEADER(
        self,
        document,
        location,
        selection,
        ensembles,
        report_dates,
        options,
    )
    _replace_visible_text(
        document,
        {
            "Сводка: ": "Основные модели: ",
            " · неопределённость: ": " · ансамбль: ",
        },
    )


def _plain_control_times_table(self, document, selection, report_dates):
    _ORIGINAL_CONTROL_TIMES(self, document, selection, report_dates)
    _replace_visible_text(
        document,
        {
            "Прогноз по контрольным срокам: через 6 ч первые 3 суток, далее через 12 ч":
                "Прогноз по времени: каждые 6 ч первые 3 суток, затем каждые 12 ч",
        },
    )


def _plain_compact_notes(document, selection, ensembles):
    _ORIGINAL_COMPACT_NOTES(document, selection, ensembles)
    _replace_visible_text(
        document,
        {
            "Ансамблевая таблица приведена отдельно на первой странице.":
                "Для ансамбля отдельно показан разброс вариантов прогноза.",
        },
    )


def _plain_ensemble_graph_page(self, document, forecast, renderer, temporary_path):
    _ORIGINAL_ENSEMBLE_GRAPH_PAGE(self, document, forecast, renderer, temporary_path)
    _replace_visible_text(
        document,
        {
            "Ансамблевая оценка — ": "Ансамбль — ",
        },
    )


apply_russian_display_names()

# Сохраняем расчётную математику, но пользовательский вывод переводим на
# простой язык без внутренних статистических терминов.
_compact_generator._short_model_name = source_display_name
_compact_generator._daily_model_metrics = consistent_daily_model_metrics
_compact_generator._daily_presentation_point = consistent_daily_presentation_point
_compact_generator._daily_temperature_text = daily_temperature_text
_compact_generator._daily_precipitation_text = daily_precipitation_metrics_text
_compact_generator._daily_wind_text = daily_wind_text
_compact_generator._daily_pressure_text = daily_pressure_text
_compact_generator._daily_agreement = daily_agreement_from_metrics
_compact_generator._build_highlights = build_highlights
_compact_generator._consensus_point = consistent_control_point
_compact_generator._detail_precipitation_text = detail_precipitation_text
_compact_generator._detail_wind_text = detail_wind_text
_compact_generator._shade_daily_hazard = shade_daily_hazard

_audit_generator._short_model_name = source_display_name
_audit_generator.ProfessionalMeteogramRenderer = DocumentMeteogramRenderer
_audit_generator.build_risk_signals = build_consistent_risk_signals
_audit_generator._daily_model_metrics = consistent_daily_model_metrics
_audit_generator._daily_presentation_point = consistent_daily_presentation_point
_audit_generator._daily_temperature_text = daily_temperature_text
_audit_generator._daily_wind_text = daily_wind_text
_audit_generator._daily_pressure_text = daily_pressure_text
_audit_generator._daily_precipitation_text_professional = daily_precipitation_text
_audit_generator._shade_daily_hazard = shade_daily_hazard

_ORIGINAL_ADD_CHART_NOTE = _audit_generator.ScientificDocumentGenerator._add_chart_note
_audit_generator.ScientificDocumentGenerator._add_chart_note = staticmethod(
    _translated_chart_note
)
_audit_generator.ScientificDocumentGenerator._add_meteogram_image = staticmethod(
    _add_meteogram_image
)

_ORIGINAL_COMPACT_HEADER = _compact_generator.ScientificDocumentGenerator._add_compact_header
_compact_generator.ScientificDocumentGenerator._add_compact_header = _plain_compact_header
_ORIGINAL_CONTROL_TIMES = _compact_generator.ScientificDocumentGenerator._add_control_times_table
_compact_generator.ScientificDocumentGenerator._add_control_times_table = _plain_control_times_table
_ORIGINAL_COMPACT_NOTES = _compact_generator.ScientificDocumentGenerator._add_compact_notes
_compact_generator.ScientificDocumentGenerator._add_compact_notes = staticmethod(
    _plain_compact_notes
)
_compact_generator.ScientificDocumentGenerator._add_daily_table = _plain_compact_daily_table
_compact_generator.ScientificDocumentGenerator._add_ensemble_summary = _plain_ensemble_summary

_audit_generator.ScientificDocumentGenerator._add_risk_cards = _plain_risk_cards
_audit_generator.ScientificDocumentGenerator._add_daily_table_professional = (
    _plain_daily_table_professional
)
_ORIGINAL_ENSEMBLE_GRAPH_PAGE = (
    _audit_generator.ScientificDocumentGenerator._add_ensemble_graph_page
)
_audit_generator.ScientificDocumentGenerator._add_ensemble_graph_page = (
    _plain_ensemble_graph_page
)

ScientificDocumentGenerator = _audit_generator.ScientificDocumentGenerator

__all__ = ["ScientificDocumentGenerator"]
