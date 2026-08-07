from __future__ import annotations

import math
import re
import statistics
from datetime import date

from docx import Document

from weather_to_docx.analysis.impact_scales import daily_precipitation_summary
from weather_to_docx.analysis.semantic_policy import strict_majority
from weather_to_docx.document import audit_generator as _audit_generator
from weather_to_docx.document import compact_generator as _compact_generator
from weather_to_docx.document.compact_generator import DailyModelMetrics
from weather_to_docx.document.source_names import source_display_name
from weather_to_docx.domain.models import ForecastPoint, ForecastSeries
from weather_to_docx.utils.files import safe_filename


def daily_precipitation_text(
    forecasts: list[ForecastSeries],
    day: date,
) -> str:
    """Describe daily precipitation without statistical jargon."""

    summaries = [
        summary
        for forecast in forecasts
        if (summary := daily_precipitation_summary(forecast, day)) is not None
    ]
    if not summaries:
        return "нет данных"

    selected_count = len(forecasts)
    available_count = len(summaries)
    totals = [summary.total_mm for summary in summaries]
    wet_count = sum(total >= 0.1 for total in totals)
    lines: list[str] = []

    if wet_count == 0:
        lines.append("без осадков")
    elif available_count == 1:
        lines.append(f"{_fmt(totals[0])} мм за сутки")
        summary = summaries[0]
        if summary.thunder:
            lines.append("возможна гроза")
        elif summary.persistent_drizzle:
            lines.append("длительная морось")
    else:
        typical = statistics.median(totals)
        low = min(totals)
        high = max(totals)
        lines.append(f"обычно около {_fmt(typical)} мм за сутки")
        if high - low >= max(0.5, typical * 0.35):
            lines.append(f"по моделям {_fmt(low)}–{_fmt(high)} мм")
        if wet_count < available_count:
            lines.append(
                f"осадки есть в {wet_count} из {available_count} моделей"
            )

        required = strict_majority(available_count)
        thunder_count = sum(summary.thunder for summary in summaries)
        drizzle_count = sum(summary.persistent_drizzle for summary in summaries)
        if 0 < thunder_count < required:
            lines.append(
                f"грозовой вариант: {thunder_count} из {available_count} моделей"
            )
        if 0 < drizzle_count < required:
            lines.append(
                "длительная морось только в "
                f"{drizzle_count} из {available_count} моделей"
            )

    missing = selected_count - available_count
    if missing > 0:
        lines.append(
            f"нет данных об осадках у {missing} из {selected_count} моделей"
        )
    return "\n".join(lines)


def daily_precipitation_metrics_text(metrics: list[DailyModelMetrics]) -> str:
    totals = [
        item.precipitation_total
        for item in metrics
        if item.precipitation_total is not None
    ]
    if not totals:
        return "нет данных"
    if max(totals) < 0.1:
        return "без осадков"
    if len(totals) == 1:
        return f"{_fmt(totals[0])} мм за сутки"

    typical = statistics.median(totals)
    low = min(totals)
    high = max(totals)
    lines = [f"обычно около {_fmt(typical)} мм за сутки"]
    if high - low >= max(0.5, typical * 0.35):
        lines.append(f"по моделям {_fmt(low)}–{_fmt(high)} мм")
    wet = sum(value >= 0.1 for value in totals)
    if wet < len(totals):
        lines.append(f"осадки есть в {wet} из {len(totals)} моделей")
    return "\n".join(lines)


def daily_wind_text(metrics: list[DailyModelMetrics]) -> str:
    winds = [item.wind_max for item in metrics if item.wind_max is not None]
    gusts = [item.gust_max for item in metrics if item.gust_max is not None]
    if not winds:
        return "нет данных"

    if len(winds) == 1:
        lines = [f"до {_fmt(winds[0])} м/с"]
    else:
        lines = [f"обычно до {_fmt(statistics.median(winds))} м/с"]

    if gusts:
        if len(gusts) == 1:
            lines.append(f"порывы до {_fmt(gusts[0])} м/с")
        else:
            typical_gust = statistics.median(gusts)
            lines.append(f"порывы обычно до {_fmt(typical_gust)} м/с")
            if max(gusts) - min(gusts) >= 3.0:
                lines.append(
                    f"по моделям {_fmt(min(gusts))}–{_fmt(max(gusts))} м/с"
                )

    if _supported(gusts, 20.0):
        lines.append("опасные порывы")
    elif _supported(gusts, 14.0):
        lines.append("сильные порывы")
    elif _supported(winds, 10.0):
        lines.append("сильный ветер")
    elif _supported(winds, 5.0):
        lines.append("ветрено")

    return "\n".join(lines)


def detail_temperature_text(points: list[ForecastPoint]) -> str:
    values = [
        value
        for point in points
        if (value := _number(point.raw("temperature_2m"))) is not None
    ]
    if not values:
        return "нет данных"
    if len(values) == 1:
        return f"{_fmt(values[0])} °C"

    typical = statistics.median(values)
    low = min(values)
    high = max(values)
    if high - low < 0.2:
        return f"около {_fmt(typical)} °C"
    return (
        f"обычно {_fmt(typical)} °C\n"
        f"по моделям {_fmt(low)}…{_fmt(high)} °C"
    )


def detail_precipitation_text(points: list[ForecastPoint]) -> str:
    available = [
        (point, value)
        for point in points
        if (value := _number(point.raw("precipitation"))) is not None
    ]
    if not available:
        return "нет данных"

    values = [value for _, value in available]
    wet = sum(value >= 0.1 for value in values)
    if wet == 0:
        return "нет"

    interval = _common_interval_hours([point for point, _ in available])
    suffix = f" за {interval:g} ч" if interval is not None else ""
    if len(values) == 1:
        return f"{_fmt(values[0])} мм{suffix}"

    typical = statistics.median(values)
    lines = [f"обычно {_fmt(typical)} мм{suffix}"]
    if max(values) - min(values) >= max(0.2, typical * 0.5):
        lines.append(f"по моделям {_fmt(min(values))}–{_fmt(max(values))} мм")
    if wet < len(values):
        lines.append(f"осадки есть в {wet} из {len(values)} моделей")
    return "\n".join(lines)


def detail_wind_text(points: list[ForecastPoint]) -> str:
    from weather_to_docx.document.consistent_controls import detail_wind_text as base

    text = base(points)
    count = sum(_number(point.raw("wind_speed_10m")) is not None for point in points)
    if count <= 1:
        return text.replace("порывы: медиана ", "порывы до ")
    return text.replace("порывы: медиана ", "порывы обычно до ")


def risk_value_text(phenomenon: str, value_text: str) -> str:
    if phenomenon == "ГРОЗА" and value_text.startswith("подтверждают "):
        return "гроза в указанный период"
    return (
        value_text.replace("медиана ", "обычно ")
        .replace("; диапазон ", "; по моделям ")
        .replace("диапазон ", "по моделям ")
    )


def risk_support_text(support_count: int, model_count: int) -> str:
    if model_count <= 1:
        return "расчёт одной модели"
    return f"подтверждают {support_count} из {model_count} моделей"


def ensemble_probability_text(
    points: list[ForecastPoint],
    threshold_token: str,
) -> tuple[str, float | None]:
    code = f"precipitation_probability_ge_{threshold_token}mm"
    candidates = [
        point.measurement(code)
        for point in points
        if point.measurement(code) is not None
    ]
    candidates = [
        item
        for item in candidates
        if item is not None and _number(item.value) is not None
    ]
    if not candidates:
        return "нет данных", None

    measurement = max(candidates, key=lambda item: float(item.value))
    probability = float(measurement.value)
    text = f"{_fmt(probability, 0)} % вариантов"
    if measurement.event_count is not None and measurement.sample_count is not None:
        text += f" ({measurement.event_count} из {measurement.sample_count})"
    return text, measurement.accumulation_hours


def ensemble_precipitation_text(points: list[ForecastPoint]) -> str:
    lines: list[str] = []
    intervals: list[float] = []
    for token, label in (("0p1", "≥0,1 мм"), ("1", "≥1 мм")):
        text, interval = ensemble_probability_text(points, token)
        lines.append(f"{label}: {text}")
        if interval is not None:
            intervals.append(float(interval))
    if intervals and max(intervals) - min(intervals) < 1e-6:
        lines.append(f"за {intervals[0]:g} ч")
    else:
        lines.append("за один шаг прогноза")
    return "\n".join(lines)


def ensemble_members_text(
    points: list[ForecastPoint],
    forecast: ForecastSeries,
) -> str:
    counts = [
        value
        for point in points
        if (value := _number(point.raw("ensemble_member_count"))) is not None
    ]
    count = int(min(counts)) if counts else forecast.source.ensemble_member_count
    expected = forecast.source.ensemble_expected_member_count
    if count is None:
        return "нет данных"
    if expected:
        return f"{count} из {expected}"
    return str(count)


def _common_interval_hours(points: list[ForecastPoint]) -> float | None:
    values: list[float] = []
    for point in points:
        measurement = point.measurement("precipitation")
        if measurement is None:
            continue
        if measurement.accumulation_hours is not None:
            values.append(float(measurement.accumulation_hours))
        elif (
            measurement.source_start_step is not None
            and measurement.source_end_step is not None
            and measurement.source_end_step > measurement.source_start_step
        ):
            values.append(
                float(measurement.source_end_step - measurement.source_start_step)
            )
    if not values:
        return None
    return values[0] if max(values) - min(values) < 1e-6 else None


def _supported(values: list[float], threshold: float) -> bool:
    if not values:
        return False
    return sum(value >= threshold for value in values) >= strict_majority(len(values))


def _number(value) -> float | None:
    try:
        number = float(value) if value is not None else None
    except (TypeError, ValueError):
        return None
    return number if number is not None and math.isfinite(number) else None


def _fmt(value: float, precision: int = 1) -> str:
    return f"{value:.{precision}f}".replace(".", ",")


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


_ORIGINAL_SOURCE_FRESHNESS = (
    _audit_generator.ScientificDocumentGenerator._add_source_freshness
)
_ORIGINAL_COMPACT_HEADER = (
    _compact_generator.ScientificDocumentGenerator._add_compact_header
)
_ORIGINAL_COMPACT_NOTES = (
    _compact_generator.ScientificDocumentGenerator._add_compact_notes
)
_ORIGINAL_GENERATE = _audit_generator.ScientificDocumentGenerator.generate


def _plain_source_freshness(document, forecast: ForecastSeries) -> None:
    _ORIGINAL_SOURCE_FRESHNESS(document, forecast)
    paragraph = document.paragraphs[-1]
    for run in paragraph.runs:
        run.text = (
            run.text.replace(
                "Цикл: цикл не передан поставщиком",
                "Расчёт модели: время запуска не указано",
            )
            .replace("Цикл:", "Расчёт модели:")
            .replace("свежие данные", "актуальные данные")
        )


def _plain_base_compact_header(
    self,
    document,
    location,
    selection,
    ensembles,
    report_dates,
    options,
) -> None:
    _ORIGINAL_COMPACT_HEADER(
        self,
        document,
        location,
        selection,
        ensembles,
        report_dates,
        options,
    )
    for paragraph in document.paragraphs:
        if "Не использованы в сводке из-за неполных данных:" not in paragraph.text:
            continue
        text = paragraph.text.replace(
            "Не использованы в сводке из-за неполных данных: ",
            "Не включены в сводку: недостаточно данных — ",
        ).replace("ключевых полей", "нужных параметров")
        text = re.sub(
            r"\((\d+) % нужных параметров\)",
            r"(доступно \1 % нужных параметров)",
            text,
        )
        _replace_paragraph_text(paragraph, text)


def _plain_base_compact_notes(document, selection, ensembles) -> None:
    _ORIGINAL_COMPACT_NOTES(document, selection, ensembles)
    for paragraph in document.paragraphs:
        if "Неполные модельные ряды в расчёт сводки не включались." in paragraph.text:
            _replace_paragraph_text(
                paragraph,
                paragraph.text.replace(
                    "Неполные модельные ряды в расчёт сводки не включались.",
                    "Модели с недостаточными данными в сводку не включались.",
                ),
            )


def _plain_base_ensemble_graph_page(
    self,
    document,
    forecast: ForecastSeries,
    renderer,
    temporary_path,
) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_after = _audit_generator.Pt(1)
    run = heading.add_run(f"Ансамбль — {source_display_name(forecast)}")
    run.bold = True
    run.font.name = "Liberation Sans"
    run.font.size = _audit_generator.Pt(12)
    self._add_source_freshness(document, forecast)

    image_path = temporary_path / (
        safe_filename(f"ensemble_{forecast.source.source_id}") + ".png"
    )
    renderer.render_ensemble(
        forecast,
        image_path,
        title=(
            f"{source_display_name(forecast)} — разброс вариантов прогноза"
        ),
    )
    self._add_meteogram_image(
        document,
        image_path,
        description=(
            f"Метеограмма вариантов ансамбля {source_display_name(forecast)}: "
            "типичное значение и разброс вариантов"
        ),
    )
    self._add_chart_note(
        document,
        "Тёмная полоса — 25–75-й процентили, светлая — 10–90-й. "
        "Вероятности осадков относятся к указанным порогам и интервалам.",
    )


def _plain_generate(
    self,
    *,
    location,
    series,
    options,
    output_path,
):
    result = _ORIGINAL_GENERATE(
        self,
        location=location,
        series=series,
        options=options,
        output_path=output_path,
    )
    document = Document(result)
    changed = False

    for paragraph in list(document.paragraphs):
        if paragraph.text.startswith("Ансамбль — это много вариантов одного расчёта."):
            _replace_paragraph_text(
                paragraph,
                "Ансамбль — это много вариантов одного расчёта. Он показывает, "
                "как может меняться прогноз. Например, 5 % вариантов — это "
                "примерно 5 из 100. Для осадков процент относится к указанной "
                "сумме за один шаг прогноза, а не за сутки.",
            )
            changed = True
        elif paragraph.text.startswith("Например, 5 % вариантов означает:"):
            paragraph._element.getparent().remove(paragraph._element)
            changed = True

    replacements = {
        "Температура\nу центральных 80 % вариантов":
            "Температура\nу 8 из 10 вариантов",
        "Порывы\n90 % вариантов не выше":
            "Порывы\nу 9 из 10 вариантов не выше",
    }
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = replacements.get(paragraph.text)
                    if new_text is not None:
                        _replace_paragraph_text(paragraph, new_text)
                        changed = True

    if changed:
        document.save(result)
    return result


# Эти подмены изменяют только представление. Расчётные медианы, квантили и
# внутренние оценки остаются прежними.
_compact_generator._detail_temperature_text = detail_temperature_text
_compact_generator.ScientificDocumentGenerator._add_compact_header = (
    _plain_base_compact_header
)
_compact_generator.ScientificDocumentGenerator._add_compact_notes = staticmethod(
    _plain_base_compact_notes
)
_audit_generator.ScientificDocumentGenerator._add_source_freshness = staticmethod(
    _plain_source_freshness
)
_audit_generator.ScientificDocumentGenerator._add_ensemble_graph_page = (
    _plain_base_ensemble_graph_page
)
_audit_generator.ScientificDocumentGenerator.generate = _plain_generate


__all__ = [
    "daily_precipitation_metrics_text",
    "daily_precipitation_text",
    "daily_wind_text",
    "detail_precipitation_text",
    "detail_temperature_text",
    "detail_wind_text",
    "ensemble_members_text",
    "ensemble_precipitation_text",
    "ensemble_probability_text",
    "risk_support_text",
    "risk_value_text",
]
