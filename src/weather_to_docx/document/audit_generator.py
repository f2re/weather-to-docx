from __future__ import annotations

import statistics
import tempfile
from datetime import UTC, date, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from weather_to_docx.analysis.consensus import (
    RiskSignal,
    build_risk_signals,
    daily_agreement,
    daily_precipitation_total,
)
from weather_to_docx.document.compact_generator import (
    AGREEMENT_HIGH,
    AGREEMENT_LOW,
    AGREEMENT_MEDIUM,
    MAX_ENSEMBLE_SYSTEMS,
    MAX_REPORT_DAYS,
    MINIMUM_MODEL_COMPLETENESS,
    _configure_compact_document,
    _daily_model_metrics,
    _daily_pressure_text,
    _daily_presentation_point,
    _daily_temperature_text,
    _daily_wind_text,
    _is_ensemble,
    _report_dates,
    _select_models,
    _set_row_height,
    _shade_daily_hazard,
    _short_model_name,
    _weekday_short,
)
from weather_to_docx.document.localized_meteogram_document import (
    ScientificDocumentGenerator as LocalizedMeteogramDocumentGenerator,
)
from weather_to_docx.document.styles import (
    DARK_BLUE,
    LIGHT_BLUE,
    WHITE,
    prevent_row_split,
    set_cell_shading,
    set_cell_text,
    set_cell_width,
    set_repeat_header_count,
    set_table_fixed_layout,
)
from weather_to_docx.document.verification import require_meteogram_docx
from weather_to_docx.document.weather_rules import weather_presentation
from weather_to_docx.domain.models import DocumentOptions, ForecastSeries, Location
from weather_to_docx.plotting.professional_meteogram import ProfessionalMeteogramRenderer
from weather_to_docx.utils.files import safe_filename

RISK_HIGH = "F8D7DA"
RISK_MEDIUM = "FFF2CC"
RISK_LOW = "EDF5FA"
RISK_CLEAR = "E2F0D9"


class ScientificDocumentGenerator(LocalizedMeteogramDocumentGenerator):
    """Оперативный документ, построенный вокруг рисков и уверенности прогноза."""

    def generate(
        self,
        *,
        location: Location,
        series: list[ForecastSeries],
        options: DocumentOptions,
        output_path: Path,
    ) -> Path:
        if not series:
            raise ValueError("Для документа не передан ни один прогностический ряд")
        if any(item.location.id != location.id for item in series):
            raise ValueError("В документ нельзя объединять прогнозы для разных координат")

        deterministic = [item for item in series if not _is_ensemble(item)]
        ensembles = [item for item in series if _is_ensemble(item)]
        selection = _select_models(
            deterministic or ensembles,
            minimum_score=MINIMUM_MODEL_COMPLETENESS,
        )
        report_dates = _report_dates(selection.reference, MAX_REPORT_DAYS)
        if not report_dates:
            raise ValueError("В прогнозе нет сроков для формирования документа")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        _configure_compact_document(document)
        self._configure_properties(document, location, options)
        self._add_compact_header(
            document,
            location,
            selection,
            ensembles,
            report_dates,
            options,
        )

        risks = build_risk_signals(
            selection.usable,
            ensembles,
            report_dates,
            maximum=3,
        )
        self._add_risk_cards(document, risks)
        self._add_daily_table_professional(document, selection.usable, report_dates)

        # Для 1–3 суток сводка и контрольные сроки используют одну страницу.
        if len(report_dates) > 3:
            document.add_page_break()
        self._add_control_times_table(document, selection, report_dates)
        self._add_compact_notes(document, selection, ensembles)

        if options.document_mode == "brief" or not options.include_meteograms:
            if options.include_ensemble_section and ensembles:
                self._add_ensemble_summary(
                    document,
                    ensembles[:MAX_ENSEMBLE_SYSTEMS],
                    report_dates,
                )
            self._add_footer(document, location)
            document.save(output_path)
            return output_path

        renderer = ProfessionalMeteogramRenderer(
            dpi=options.meteogram_dpi,
            smoothing=options.meteogram_smoothing,
        )
        with tempfile.TemporaryDirectory(prefix="weather-to-docx-audit-") as temporary:
            temporary_path = Path(temporary)
            for forecast in selection.usable:
                if _is_ensemble(forecast):
                    continue
                document.add_page_break()
                self._add_model_graph_page(
                    document,
                    forecast,
                    renderer,
                    temporary_path,
                )
                if options.document_mode == "full":
                    document.add_page_break()
                    self._add_model_daily_table(document, forecast, report_dates)

            if options.include_ensemble_section:
                for forecast in ensembles[:MAX_ENSEMBLE_SYSTEMS]:
                    document.add_page_break()
                    self._add_ensemble_graph_page(
                        document,
                        forecast,
                        renderer,
                        temporary_path,
                    )
                    if options.document_mode == "full":
                        document.add_page_break()
                        self._add_ensemble_summary(document, [forecast], report_dates)

            self._add_footer(document, location)
            document.save(output_path)
            require_meteogram_docx(output_path)
        return output_path

    def _add_risk_cards(
        self,
        document: Document,
        risks: list[RiskSignal],
    ) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(2)
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run("Ключевые риски")
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(10.5)

        if not risks:
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            set_table_fixed_layout(table)
            set_cell_shading(table.cell(0, 0), RISK_CLEAR)
            set_cell_text(
                table.cell(0, 0),
                "Существенные погодные риски не подтверждены выбранными моделями и ансамблем.",
                size=8,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
            return

        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(table)
        widths = (43, 65, 71, 88)
        headers = ("Явление", "Время", "Подтверждение", "Оценка")
        for cell, text, width in zip(table.rows[0].cells, headers, widths, strict=True):
            set_cell_text(cell, text, size=7.2, bold=True)
            set_cell_shading(cell, DARK_BLUE)
            self._set_text_color(cell, WHITE)
            set_cell_width(cell, width)
        set_repeat_header_count(table, 1)

        for signal in risks:
            row = table.add_row()
            prevent_row_split(row)
            fill = (
                RISK_HIGH
                if signal.severity >= 85
                else RISK_MEDIUM
                if signal.severity >= 70
                else RISK_LOW
            )
            for cell, width in zip(row.cells, widths, strict=True):
                set_cell_width(cell, width)
                set_cell_shading(cell, fill)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            probability = (
                f"\nансамбль осадков: {signal.ensemble_probability:.0f} %"
                if signal.ensemble_probability is not None
                else ""
            )
            peak = (
                f"\nмаксимум около {signal.peak_local:%H:%M}"
                if signal.peak_local is not None
                else ""
            )
            values = (
                signal.phenomenon,
                signal.time_text + peak,
                f"{signal.support_text}{probability}",
                f"{signal.scenario}\n{signal.value_text}\nуверенность: {signal.confidence}",
            )
            for index, (cell, text) in enumerate(zip(row.cells, values, strict=True)):
                set_cell_text(
                    cell,
                    text,
                    size=7.2,
                    bold=index in {0, 3},
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                )

    def _add_daily_table_professional(
        self,
        document: Document,
        forecasts: list[ForecastSeries],
        report_dates: list[date],
    ) -> None:
        show_agreement = len(forecasts) > 1
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(3)
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run("Прогноз по дням")
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(10)

        headers = ["Дата", "Погода", "Температура", "Осадки", "Ветер", "Давление"]
        widths = [24, 48, 34, 47, 56, 31]
        if show_agreement:
            headers.append("Уверенность")
            widths.append(37)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(table)
        for cell, text, width in zip(table.rows[0].cells, headers, widths, strict=True):
            set_cell_text(cell, text, size=7.2, bold=True)
            set_cell_shading(cell, DARK_BLUE)
            self._set_text_color(cell, WHITE)
            set_cell_width(cell, width)
        set_repeat_header_count(table, 1)

        agreement_notes: list[str] = []
        for day in report_dates:
            metrics = [
                metric
                for forecast in forecasts
                if (metric := _daily_model_metrics(forecast, day)) is not None
            ]
            if not metrics:
                continue
            representative = _daily_presentation_point(metrics, day)
            row = table.add_row()
            prevent_row_split(row)
            _set_row_height(row, 8)
            for cell, width in zip(row.cells, widths, strict=True):
                set_cell_width(cell, width)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            values = [
                f"{day:%d.%m}\n{_weekday_short(day)}",
                None,
                _daily_temperature_text(metrics),
                _daily_precipitation_text_professional(forecasts, day),
                _daily_wind_text(metrics),
                _daily_pressure_text(metrics),
            ]
            set_cell_text(row.cells[0], values[0], size=7.2, bold=True)
            presentation = weather_presentation(representative)
            self._set_compact_icon_cell(
                row.cells[1],
                presentation.icon_key,
                presentation.description,
            )
            for index in range(2, 6):
                set_cell_text(row.cells[index], values[index], size=7)

            if show_agreement:
                agreement = daily_agreement(forecasts, day)
                label = agreement.overall if agreement else "нет оценки"
                fill = {
                    "высокая": AGREEMENT_HIGH,
                    "средняя": AGREEMENT_MEDIUM,
                    "низкая": AGREEMENT_LOW,
                }.get(label, LIGHT_BLUE)
                set_cell_text(row.cells[6], label, size=7, bold=True)
                set_cell_shading(row.cells[6], fill)
                if agreement and agreement.overall != "высокая":
                    agreement_notes.append(f"{day:%d.%m}: {agreement.note}.")
            _shade_daily_hazard(row.cells, metrics)

        if agreement_notes:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run("Расхождения моделей: " + " ".join(agreement_notes))
            run.font.name = "Liberation Sans"
            run.font.size = Pt(6.8)
            run.font.italic = True

    def _add_model_graph_page(
        self,
        document: Document,
        forecast: ForecastSeries,
        renderer: ProfessionalMeteogramRenderer,
        temporary_path: Path,
    ) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_after = Pt(1)
        run = heading.add_run(f"Метеограмма модели {_short_model_name(forecast)}")
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(12)
        self._add_source_freshness(document, forecast)

        image_path = temporary_path / (
            safe_filename(f"model_{forecast.source.source_id}") + ".png"
        )
        renderer.render_deterministic(
            forecast,
            image_path,
            title=f"{_short_model_name(forecast)} — весь срок прогноза",
        )
        self._add_meteogram_image(
            document,
            image_path,
            description=f"Метеограмма модели {_short_model_name(forecast)} на весь срок",
        )
        self._add_chart_note(
            document,
            "Серые полосы обозначают ночь. Осадки показаны за исходный интервал; "
            "стрелки внизу показывают направление ветра.",
        )

    def _add_ensemble_graph_page(
        self,
        document: Document,
        forecast: ForecastSeries,
        renderer: ProfessionalMeteogramRenderer,
        temporary_path: Path,
    ) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_after = Pt(1)
        run = heading.add_run(f"Ансамблевая оценка — {_short_model_name(forecast)}")
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(12)
        self._add_source_freshness(document, forecast)

        image_path = temporary_path / (
            safe_filename(f"ensemble_{forecast.source.source_id}") + ".png"
        )
        renderer.render_ensemble(
            forecast,
            image_path,
            title=f"{_short_model_name(forecast)} — неопределённость прогноза",
        )
        self._add_meteogram_image(
            document,
            image_path,
            description=(
                f"Ансамблевая метеограмма {_short_model_name(forecast)}: "
                "медиана, межквартильный и 10–90-процентильный диапазоны"
            ),
        )
        self._add_chart_note(
            document,
            "Тёмная полоса — 25–75-й процентили, светлая — 10–90-й. "
            "Вероятности осадков относятся к указанным порогам и интервалам.",
        )

    @staticmethod
    def _add_source_freshness(document: Document, forecast: ForecastSeries) -> None:
        source = forecast.source
        cycle = (
            source.cycle_time_utc.astimezone(UTC).strftime("%d.%m.%Y %H:%M UTC")
            if source.cycle_time_utc
            else "цикл не передан поставщиком"
        )
        retrieved = source.retrieved_at_utc.astimezone(UTC).strftime("%d.%m.%Y %H:%M UTC")
        age_hours = max(
            0.0,
            (datetime.now(UTC) - source.retrieved_at_utc.astimezone(UTC)).total_seconds() / 3600,
        )
        state = "свежие данные" if age_hours <= 12 else "данные старше 12 часов"
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(f"Цикл: {cycle} · получено: {retrieved} · {state}")
        run.font.name = "Liberation Sans"
        run.font.size = Pt(7)
        run.font.bold = age_hours > 12


def _daily_precipitation_text_professional(
    forecasts: list[ForecastSeries],
    day: date,
) -> str:
    totals = []
    for forecast in forecasts:
        points = [
            point
            for point in forecast.points
            if point.valid_time_local.date() == day
        ]
        total = daily_precipitation_total(points)
        if total is not None:
            totals.append(total)
    if not totals:
        return "нет данных"
    wet_count = sum(total >= 0.1 for total in totals)
    if max(totals) < 0.1:
        return "без осадков"
    low = min(totals)
    high = max(totals)
    amount = (
        f"{statistics.median(totals):.1f} мм"
        if high - low < 0.05
        else f"{low:.1f}–{high:.1f} мм"
    ).replace(".", ",")
    return f"{amount}\nосадки: {wet_count}/{len(totals)} моделей"
