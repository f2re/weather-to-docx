from __future__ import annotations

from docx import Document

from weather_to_docx.document.meteogram_document import (
    ScientificDocumentGenerator as MeteogramDocumentGenerator,
)
from weather_to_docx.document.russian_labels import (
    RUSSIAN_MODEL_NAMES,
    visible_timezone_label,
)
from weather_to_docx.domain.models import (
    DocumentOptions,
    ForecastSeries,
    Location,
)


class ScientificDocumentGenerator(MeteogramDocumentGenerator):
    """Генератор без англоязычных служебных подписей в видимой части."""

    def _add_compact_header(
        self,
        document: Document,
        location: Location,
        selection,
        ensembles: list[ForecastSeries],
        report_dates,
        options: DocumentOptions,
    ) -> None:
        super()._add_compact_header(
            document,
            location,
            selection,
            ensembles,
            report_dates,
            options,
        )
        local_time = selection.reference.points[0].valid_time_local
        replacements = {
            location.timezone: visible_timezone_label(local_time),
            **RUSSIAN_MODEL_NAMES,
        }
        _replace_visible_text(document, replacements)

    @staticmethod
    def _add_chart_note(document: Document, text: str) -> None:
        text = text.replace(
            "shape-preserving методом PCHIP",
            "формосохраняющим методом PCHIP",
        )
        text = text.replace(
            "shape-preserving",
            "формосохраняющим",
        )
        MeteogramDocumentGenerator._add_chart_note(document, text)


def _replace_visible_text(
    document: Document,
    replacements: dict[str, str],
) -> None:
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    for run in paragraph.runs:
        updated = run.text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        run.text = updated
