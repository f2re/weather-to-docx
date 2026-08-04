from __future__ import annotations

import math
import statistics
from collections import Counter
from dataclasses import dataclass
from datetime import UTC, date, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from weather_to_docx.document.generator import DocumentGenerator
from weather_to_docx.document.styles import (
    DANGER,
    DARK_BLUE,
    LIGHT_BLUE,
    WARNING,
    WHITE,
    configure_document,
    prevent_row_split,
    set_cell_shading,
    set_cell_text,
    set_cell_width,
    set_repeat_header_count,
    set_table_fixed_layout,
)
from weather_to_docx.document.weather_rules import weather_presentation
from weather_to_docx.domain.models import (
    DocumentOptions,
    ForecastPoint,
    ForecastSeries,
    ForecastValue,
    Location,
    SourceKind,
)
from weather_to_docx.utils.meteorology import wind_rumb

AGREEMENT_HIGH = "E2F0D9"
AGREEMENT_MEDIUM = "FFF2CC"
AGREEMENT_LOW = "FCE4D6"

MAX_REPORT_DAYS = 7
MAX_DETAIL_ROWS = 18
MINIMUM_MODEL_COMPLETENESS = 0.55
MAX_ENSEMBLE_SYSTEMS = 1
MAX_HIGHLIGHTS = 4
DETAIL_INTERVAL_HOURS = 6
EXTENDED_DETAIL_INTERVAL_HOURS = 12
DETAIL_SWITCH_HOUR = 72

KEY_PARAMETERS = (
    "temperature_2m",
    "precipitation",
    "wind_speed_10m",
    "wind_gusts_10m",
    "pressure_msl",
    "cloud_cover",
    "relative_humidity_2m",
)
REQUIRED_PARAMETERS = (
    "temperature_2m",
    "wind_speed_10m",
    "pressure_msl",
)

MODEL_SHORT_NAMES = {
    "NOAA GFS 0.25°": "GFS",
    "ECMWF IFS 0.25° Open Data": "ECMWF IFS",
    "ECMWF AIFS 0.25° Single": "ECMWF AIFS",
    "DWD ICON Global": "ICON",
    "ECCC GEM Global (GDPS)": "GDPS",
    "NOAA GEFS 0.25°": "GEFS",
    "NOAA GEFS 0.5°": "GEFS",
    "ECMWF IFS Ensemble 0.25°": "ECMWF ENS",
    "ECMWF AIFS Ensemble 0.25°": "AIFS ENS",
    "DWD ICON Global EPS": "ICON-EPS",
    "ECCC Global Ensemble Prediction System": "GEPS",
}

WEATHER_SEVERITY = {
    0: 0,
    1: 1,
    2: 2,
    3: 3,
    45: 6,
    48: 7,
    51: 4,
    53: 5,
    55: 6,
    56: 7,
    57: 8,
    61: 5,
    63: 7,
    65: 9,
    66: 8,
    67: 10,
    71: 6,
    73: 8,
    75: 10,
    77: 7,
    80: 6,
    81: 8,
    82: 10,
    85: 7,
    86: 9,
    95: 11,
    96: 12,
    99: 13,
}


@dataclass(frozen=True, slots=True)
class ModelSelection:
    usable: list[ForecastSeries]
    excluded: list[tuple[ForecastSeries, float]]
    reference: ForecastSeries


@dataclass(frozen=True, slots=True)
class DailyModelMetrics:
    source: ForecastSeries
    weather_code: int
    temperature_min: float | None
    temperature_max: float | None
    precipitation_total: float | None
    wind_max: float | None
    gust_max: float | None
    pressure_min: float | None
    pressure_max: float | None


class ScientificDocumentGenerator(DocumentGenerator):
    """Короткий операторский DOCX: одна сводка и одна таблица сроков.

    Модели не размножаются по отдельным разделам. Они используются как
    независимые оценки одного прогноза, а неполные ряды автоматически
    исключаются из сводки. Ансамбль остаётся отдельной компактной таблицей.
    """

    def generate(
        self,
        *,
        location: Location,
        series: list[ForecastSeries],
        options: DocumentOptions,
        output_path: Path,
    ) -> Path:
        if not series:
            raise ValueError("Для документа не передан ни один прогностический ряд")
        if any(item.location.id != location.id for item in series):
            raise ValueError(
                "В документ нельзя объединять прогнозы для разных координат"
            )

        deterministic = [item for item in series if not _is_ensemble(item)]
        ensembles = [item for item in series if _is_ensemble(item)]
        candidates = deterministic or ensembles
        selection = _select_models(
            candidates,
            minimum_score=MINIMUM_MODEL_COMPLETENESS,
        )
        report_dates = _report_dates(selection.reference, MAX_REPORT_DAYS)
        if not report_dates:
            raise ValueError("В прогнозе нет сроков для формирования документа")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        _configure_compact_document(document)
        self._configure_properties(document, location, options)
        self._add_compact_header(
            document,
            location,
            selection,
            ensembles,
            report_dates,
            options,
        )
        self._add_highlights(document, selection, report_dates)
        self._add_daily_table(document, selection, report_dates)
        self._add_ensemble_summary(
            document,
            ensembles[:MAX_ENSEMBLE_SYSTEMS],
            report_dates,
        )

        document.add_page_break()
        self._add_control_times_table(document, selection, report_dates)
        self._add_compact_notes(document, selection, ensembles)
        self._add_footer(document, location)
        document.save(output_path)
        return output_path

    def _add_compact_header(
        self,
        document: Document,
        location: Location,
        selection: ModelSelection,
        ensembles: list[ForecastSeries],
        report_dates: list[date],
        options: DocumentOptions,
    ) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(0)
        run = title.add_run(options.title)
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(15)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(2)
        run = subtitle.add_run(location.name)
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(11)

        period = (
            f"{report_dates[0]:%d.%m.%Y}–{report_dates[-1]:%d.%m.%Y}"
            if len(report_dates) > 1
            else f"{report_dates[0]:%d.%m.%Y}"
        )
        model_names = ", ".join(_short_model_name(item) for item in selection.usable)
        ensemble_names = ", ".join(_short_model_name(item) for item in ensembles)
        local_generated = datetime.now(UTC).astimezone(
            selection.reference.points[0].valid_time_local.tzinfo
        )
        metadata = document.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata.paragraph_format.space_after = Pt(2)
        text = (
            f"Период: {period} · координаты {location.latitude:.4f}, "
            f"{location.longitude:.4f} · местное время {location.timezone} · "
            f"обновлено {local_generated:%d.%m.%Y %H:%M}"
        )
        run = metadata.add_run(text)
        run.font.name = "Liberation Sans"
        run.font.size = Pt(7.2)

        sources = document.add_paragraph()
        sources.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sources.paragraph_format.space_after = Pt(3)
        source_text = f"Сводка: {model_names}"
        if ensemble_names:
            source_text += f" · неопределённость: {ensemble_names}"
        run = sources.add_run(source_text)
        run.font.name = "Liberation Sans"
        run.font.size = Pt(7.2)
        run.bold = True

        if selection.excluded:
            excluded = ", ".join(
                f"{_short_model_name(item)} ({score * 100:.0f} % ключевых полей)"
                for item, score in selection.excluded
            )
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(
                "Не использованы в сводке из-за неполных данных: " + excluded
            )
            run.font.name = "Liberation Sans"
            run.font.size = Pt(7)
            run.bold = True

    def _add_highlights(
        self,
        document: Document,
        selection: ModelSelection,
        report_dates: list[date],
    ) -> None:
        highlights = _build_highlights(
            selection.usable,
            report_dates,
            MAX_HIGHLIGHTS,
        )
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(3)
        heading = paragraph.add_run("Главное: ")
        heading.bold = True
        heading.font.name = "Liberation Sans"
        heading.font.size = Pt(8)
        text = (
            " • ".join(highlights)
            if highlights
            else "существенных погодных рисков по выбранным моделям не выделено."
        )
        run = paragraph.add_run(text)
        run.font.name = "Liberation Sans"
        run.font.size = Pt(8)

    def _add_daily_table(
        self,
        document: Document,
        selection: ModelSelection,
        report_dates: list[date],
    ) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(1)
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run("Прогноз по дням")
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(10)

        headers = (
            "Дата",
            "Погода",
            "Температура",
            "Осадки",
            "Ветер",
            "Давление",
            "Согласованность",
        )
        widths = (24, 45, 31, 42, 52, 30, 35)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(table)
        for cell, text, width in zip(table.rows[0].cells, headers, widths, strict=True):
            set_cell_text(cell, text, size=7.2, bold=True)
            set_cell_shading(cell, DARK_BLUE)
            self._set_text_color(cell, WHITE)
            set_cell_width(cell, width)
        _set_row_height(table.rows[0], 6)
        set_repeat_header_count(table, 1)

        for day in report_dates:
            metrics = [
                metric
                for forecast in selection.usable
                if (metric := _daily_model_metrics(forecast, day)) is not None
            ]
            if not metrics:
                continue
            representative = _daily_presentation_point(metrics, day)
            row = table.add_row()
            prevent_row_split(row)
            _set_row_height(row, 8)
            for cell, width in zip(row.cells, widths, strict=True):
                set_cell_width(cell, width)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            agreement, agreement_fill = _daily_agreement(metrics)
            set_cell_text(
                row.cells[0],
                f"{day:%d.%m}\n{_weekday_short(day)}",
                size=7.2,
                bold=True,
            )
            presentation = weather_presentation(representative)
            self._set_compact_icon_cell(
                row.cells[1],
                presentation.icon_key,
                presentation.description,
            )
            set_cell_text(row.cells[2], _daily_temperature_text(metrics), size=7)
            set_cell_text(row.cells[3], _daily_precipitation_text(metrics), size=7)
            set_cell_text(row.cells[4], _daily_wind_text(metrics), size=7)
            set_cell_text(row.cells[5], _daily_pressure_text(metrics), size=7)
            set_cell_text(row.cells[6], agreement, size=7, bold=True)
            _shade_daily_hazard(row.cells, metrics)
            set_cell_shading(row.cells[6], agreement_fill)

    def _add_ensemble_summary(
        self,
        document: Document,
        ensembles: list[ForecastSeries],
        report_dates: list[date],
    ) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(3)
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run("Неопределённость прогноза")
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(10)

        if not ensembles:
            paragraph = document.add_paragraph(
                "Ансамблевая система не выбрана; вероятностная оценка не рассчитана."
            )
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.name = "Liberation Sans"
                run.font.size = Pt(7.2)
            return

        headers = (
            "Дата",
            "Ансамбль",
            "Температура, 10–90 %",
            "Осадки ≥0,1 мм",
            "Осадки ≥1 мм",
            "Порывы, 90 %",
            "Члены",
        )
        widths = (23, 36, 46, 42, 38, 37, 35)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(table)
        for cell, text, width in zip(table.rows[0].cells, headers, widths, strict=True):
            set_cell_text(cell, text, size=6.8, bold=True)
            set_cell_shading(cell, DARK_BLUE)
            self._set_text_color(cell, WHITE)
            set_cell_width(cell, width)
        _set_row_height(table.rows[0], 6)
        set_repeat_header_count(table, 1)

        row_count = 0
        for forecast in ensembles:
            for day in report_dates:
                points = [
                    point
                    for point in forecast.points
                    if point.valid_time_local.date() == day
                ]
                if not points:
                    continue
                row = table.add_row()
                prevent_row_split(row)
                _set_row_height(row, 5.8)
                for cell, width in zip(row.cells, widths, strict=True):
                    set_cell_width(cell, width)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_shading(
                        cell,
                        LIGHT_BLUE if row_count % 2 == 0 else WHITE,
                    )
                row_count += 1
                noon = min(points, key=lambda point: abs(point.valid_time_local.hour - 12))
                set_cell_text(row.cells[0], day.strftime("%d.%m"), size=6.6, bold=True)
                set_cell_text(row.cells[1], _short_model_name(forecast), size=6.4)
                set_cell_text(row.cells[2], _ensemble_temperature_text(noon), size=6.4)
                set_cell_text(
                    row.cells[3],
                    _ensemble_probability_text(points, "0p1"),
                    size=6.4,
                )
                set_cell_text(
                    row.cells[4],
                    _ensemble_probability_text(points, "1"),
                    size=6.4,
                )
                set_cell_text(row.cells[5], _ensemble_gust_text(points), size=6.4)
                set_cell_text(
                    row.cells[6],
                    _ensemble_members_text(points, forecast),
                    size=6.4,
                )

        note = document.add_paragraph()
        note.paragraph_format.space_before = Pt(1)
        note.paragraph_format.space_after = Pt(0)
        run = note.add_run(
            "Диапазон 10–90 % охватывает центральные 80 % вариантов ансамбля; "
            "вероятность осадков — максимальная доля членов за сутки."
        )
        run.font.name = "Liberation Sans"
        run.font.size = Pt(6.6)

    def _add_control_times_table(
        self,
        document: Document,
        selection: ModelSelection,
        report_dates: list[date],
    ) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_after = Pt(1)
        run = heading.add_run(
            "Прогноз по контрольным срокам: через 6 ч первые 3 суток, далее через 12 ч"
        )
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(10)

        headers = (
            "Дата и время",
            "Погода",
            "Температура",
            "Осадки",
            "Ветер",
            "Влажность",
            "Давление",
        )
        widths = (32, 48, 37, 45, 55, 29, 28)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(table)
        for cell, text, width in zip(table.rows[0].cells, headers, widths, strict=True):
            set_cell_text(cell, text, size=7, bold=True)
            set_cell_shading(cell, DARK_BLUE)
            self._set_text_color(cell, WHITE)
            set_cell_width(cell, width)
        _set_row_height(table.rows[0], 6)
        set_repeat_header_count(table, 1)

        allowed_dates = set(report_dates)
        reference_points = _control_points(selection.reference.points, allowed_dates)
        indexes = {
            forecast.source.source_id: {
                point.valid_time_utc: point for point in forecast.points
            }
            for forecast in selection.usable
        }
        previous_date = None
        for reference in reference_points:
            points = [
                index[reference.valid_time_utc]
                for index in indexes.values()
                if reference.valid_time_utc in index
            ]
            if not points:
                continue
            consensus = _consensus_point(points, reference)
            row = table.add_row()
            prevent_row_split(row)
            _set_row_height(row, 7)
            fill = (
                LIGHT_BLUE
                if reference.valid_time_local.date() != previous_date
                else WHITE
            )
            previous_date = reference.valid_time_local.date()
            for cell, width in zip(row.cells, widths, strict=True):
                set_cell_width(cell, width)
                set_cell_shading(cell, fill)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            set_cell_text(
                row.cells[0],
                reference.valid_time_local.strftime("%d.%m\n%H:%M"),
                size=7,
                bold=True,
            )
            presentation = weather_presentation(consensus)
            self._set_compact_icon_cell(
                row.cells[1],
                presentation.icon_key,
                presentation.description,
            )
            set_cell_text(row.cells[2], _detail_temperature_text(points), size=6.8)
            set_cell_text(row.cells[3], _detail_precipitation_text(points), size=6.8)
            set_cell_text(row.cells[4], _detail_wind_text(points), size=6.8)
            set_cell_text(row.cells[5], _detail_humidity_text(points), size=6.8)
            set_cell_text(row.cells[6], _detail_pressure_text(points), size=6.8)
            self._apply_hazard_shading(row.cells, consensus)

    def _set_compact_icon_cell(
        self,
        cell,
        icon_key: str,
        description: str,
    ) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.9
        run = paragraph.add_run()
        shape = run.add_picture(str(self.icons.render(icon_key)), width=Mm(5))
        shape._inline.docPr.set("descr", f"Погода: {description}")
        shape._inline.docPr.set("title", description)
        description_run = paragraph.add_run(f"\n{description}")
        description_run.font.name = "Liberation Sans"
        description_run.font.size = Pt(6.1)

    @staticmethod
    def _add_compact_notes(
        document: Document,
        selection: ModelSelection,
        ensembles: list[ForecastSeries],
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(0)
        used = ", ".join(_short_model_name(item) for item in selection.usable)
        text = f"Сводные значения рассчитаны по моделям: {used}."
        if ensembles:
            text += " Ансамблевая таблица приведена отдельно на первой странице."
        if selection.excluded:
            text += " Неполные модельные ряды в расчёт сводки не включались."
        run = paragraph.add_run(text)
        run.font.name = "Liberation Sans"
        run.font.size = Pt(6.6)


def _configure_compact_document(document: Document) -> None:
    configure_document(document, "A4")
    section = document.sections[0]
    section.top_margin = Mm(7)
    section.bottom_margin = Mm(7)
    section.left_margin = Mm(7)
    section.right_margin = Mm(7)
    normal = document.styles["Normal"]
    normal.font.name = "Liberation Sans"
    normal.font.size = Pt(8)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = 0.95


def _set_row_height(row, millimetres: float) -> None:
    row.height = Mm(millimetres)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _is_ensemble(forecast: ForecastSeries) -> bool:
    return (
        forecast.source.source_kind == SourceKind.ENSEMBLE
        or forecast.source.ensemble_member_count is not None
        or any(
            token in f"{forecast.source.source_id} {forecast.source.model}".lower()
            for token in ("ensemble", "gefs", "geps", "eps")
        )
    )


def _select_models(
    forecasts: list[ForecastSeries],
    *,
    minimum_score: float,
) -> ModelSelection:
    scored = sorted(
        ((forecast, _model_completeness(forecast)) for forecast in forecasts),
        key=lambda item: (item[1], len(item[0].points)),
        reverse=True,
    )
    if not scored:
        raise ValueError("Не передано ни одной модели для сводного прогноза")

    usable = [forecast for forecast, score in scored if score >= minimum_score]
    excluded = [(forecast, score) for forecast, score in scored if score < minimum_score]
    if not usable:
        best, best_score = scored[0]
        if best_score < 0.35:
            missing = ", ".join(_missing_key_parameters(best))
            raise ValueError(
                "Ни одна модель не содержит достаточного набора ключевых полей. "
                f"Лучшая модель {_short_model_name(best)}: {best_score * 100:.0f} %. "
                f"Отсутствуют: {missing or 'ключевые параметры'}."
            )
        usable = [best]
        excluded = [(forecast, score) for forecast, score in scored[1:]]
    reference = max(
        usable,
        key=lambda forecast: (_model_completeness(forecast), len(forecast.points)),
    )
    return ModelSelection(usable=usable, excluded=excluded, reference=reference)


def _model_completeness(forecast: ForecastSeries) -> float:
    if not forecast.points:
        return 0.0
    weights = {
        "temperature_2m": 2.0,
        "precipitation": 1.2,
        "wind_speed_10m": 1.5,
        "wind_gusts_10m": 0.8,
        "pressure_msl": 1.2,
        "cloud_cover": 0.7,
        "relative_humidity_2m": 0.6,
    }
    total_weight = sum(weights.values())
    score = sum(
        weights[code] * _parameter_coverage(forecast, code)
        for code in KEY_PARAMETERS
    ) / total_weight
    required = statistics.fmean(
        _parameter_coverage(forecast, code) for code in REQUIRED_PARAMETERS
    )
    if required < 0.5:
        score *= 0.45
    return score


def _parameter_coverage(forecast: ForecastSeries, code: str) -> float:
    if not forecast.points:
        return 0.0
    available = sum(_numeric(point, code) is not None for point in forecast.points)
    return available / len(forecast.points)


def _missing_key_parameters(forecast: ForecastSeries) -> list[str]:
    names = {
        "temperature_2m": "температура",
        "precipitation": "осадки",
        "wind_speed_10m": "ветер",
        "wind_gusts_10m": "порывы",
        "pressure_msl": "давление",
        "cloud_cover": "облачность",
        "relative_humidity_2m": "влажность",
    }
    return [
        names[code]
        for code in KEY_PARAMETERS
        if _parameter_coverage(forecast, code) < 0.5
    ]


def _report_dates(reference: ForecastSeries, maximum_days: int) -> list[date]:
    values = sorted({point.valid_time_local.date() for point in reference.points})
    return values[:maximum_days]


def _daily_model_metrics(
    forecast: ForecastSeries,
    day: date,
) -> DailyModelMetrics | None:
    points = [
        point
        for point in forecast.points
        if point.valid_time_local.date() == day
    ]
    if not points:
        return None
    temperatures = _values(points, "temperature_2m")
    precipitation = _values(points, "precipitation")
    winds = _values(points, "wind_speed_10m")
    gusts = _values(points, "wind_gusts_10m")
    pressure = _values(points, "pressure_msl")
    weather_codes = [weather_presentation(point).code for point in points]
    noon = min(points, key=lambda point: abs(point.valid_time_local.hour - 12))
    weather_code = weather_presentation(noon).code
    severe_codes = [code for code in weather_codes if WEATHER_SEVERITY.get(code, 3) >= 5]
    if severe_codes and (sum(precipitation) if precipitation else 0) >= 0.1:
        weather_code = max(
            severe_codes,
            key=lambda code: WEATHER_SEVERITY.get(code, 3),
        )
    return DailyModelMetrics(
        source=forecast,
        weather_code=weather_code,
        temperature_min=min(temperatures) if temperatures else None,
        temperature_max=max(temperatures) if temperatures else None,
        precipitation_total=(
            sum(max(0.0, value) for value in precipitation)
            if precipitation
            else None
        ),
        wind_max=max(winds) if winds else None,
        gust_max=max(gusts) if gusts else None,
        pressure_min=min(pressure) if pressure else None,
        pressure_max=max(pressure) if pressure else None,
    )


def _daily_presentation_point(
    metrics: list[DailyModelMetrics],
    day: date,
) -> ForecastPoint:
    codes = [metric.weather_code for metric in metrics]
    counts = Counter(codes)
    code = max(
        codes,
        key=lambda item: (counts[item], WEATHER_SEVERITY.get(item, 3)),
    )
    local_time = min(
        (
            point.valid_time_local
            for metric in metrics
            for point in metric.source.points
            if point.valid_time_local.date() == day
        ),
        key=lambda value: abs(value.hour - 12),
    )
    return ForecastPoint(
        valid_time_utc=local_time.astimezone(UTC),
        valid_time_local=local_time,
        weather_code=code,
        is_day=True,
        values={},
    )


def _daily_temperature_text(metrics: list[DailyModelMetrics]) -> str:
    lows = [value for item in metrics if (value := item.temperature_min) is not None]
    highs = [value for item in metrics if (value := item.temperature_max) is not None]
    if not lows or not highs:
        return "нет данных"
    return f"{_fmt(statistics.median(lows))}…{_fmt(statistics.median(highs))} °C"


def _daily_precipitation_text(metrics: list[DailyModelMetrics]) -> str:
    totals = [
        value
        for item in metrics
        if (value := item.precipitation_total) is not None
    ]
    if not totals:
        return "нет данных"
    wet = sum(value >= 0.1 for value in totals)
    if max(totals) < 0.1:
        return "без осадков"
    if max(totals) - min(totals) < 0.05:
        amount = f"{_fmt(statistics.median(totals))} мм"
    else:
        amount = f"{_fmt(min(totals))}–{_fmt(max(totals))} мм"
    return f"{amount}\nосадки: {wet}/{len(totals)} моделей"


def _daily_wind_text(metrics: list[DailyModelMetrics]) -> str:
    winds = [value for item in metrics if (value := item.wind_max) is not None]
    gusts = [value for item in metrics if (value := item.gust_max) is not None]
    if not winds:
        return "нет данных"
    text = f"до {_fmt(statistics.median(winds))} м/с"
    if gusts:
        text += f"\nпорывы до {_fmt(max(gusts))} м/с"
    return text


def _daily_pressure_text(metrics: list[DailyModelMetrics]) -> str:
    lows = [value for item in metrics if (value := item.pressure_min) is not None]
    highs = [value for item in metrics if (value := item.pressure_max) is not None]
    if not lows or not highs:
        return "нет данных"
    return f"{_fmt(min(lows), 0)}–{_fmt(max(highs), 0)} гПа"


def _daily_agreement(metrics: list[DailyModelMetrics]) -> tuple[str, str]:
    if len(metrics) <= 1:
        return "одна модель", AGREEMENT_MEDIUM
    highs = [value for item in metrics if (value := item.temperature_max) is not None]
    wet = [
        item.precipitation_total is not None and item.precipitation_total >= 0.1
        for item in metrics
    ]
    codes = [item.weather_code for item in metrics]
    weather_ratio = max(Counter(codes).values()) / len(codes)
    temperature_spread = max(highs) - min(highs) if len(highs) >= 2 else 0.0
    precipitation_agrees = all(wet) or not any(wet)
    if weather_ratio >= 0.67 and temperature_spread <= 2.5 and precipitation_agrees:
        return "высокая", AGREEMENT_HIGH
    if weather_ratio >= 0.5 and temperature_spread <= 4.5:
        return "средняя", AGREEMENT_MEDIUM
    return "низкая", AGREEMENT_LOW


def _shade_daily_hazard(cells, metrics: list[DailyModelMetrics]) -> None:
    precipitation = max(
        (item.precipitation_total or 0.0 for item in metrics),
        default=0.0,
    )
    gust = max((item.gust_max or 0.0 for item in metrics), default=0.0)
    highs = [value for item in metrics if (value := item.temperature_max) is not None]
    lows = [value for item in metrics if (value := item.temperature_min) is not None]
    dangerous = precipitation >= 10 or gust >= 20
    warning = (
        precipitation >= 3
        or gust >= 14
        or (highs and max(highs) >= 35)
        or (lows and min(lows) <= -25)
    )
    if dangerous:
        for cell in cells:
            set_cell_shading(cell, DANGER)
    elif warning:
        for cell in cells:
            set_cell_shading(cell, WARNING)


def _build_highlights(
    forecasts: list[ForecastSeries],
    report_dates: list[date],
    maximum: int,
) -> list[str]:
    highlights: list[tuple[int, str]] = []
    for day in report_dates:
        metrics = [
            metric
            for forecast in forecasts
            if (metric := _daily_model_metrics(forecast, day)) is not None
        ]
        if not metrics:
            continue
        gust = max((item.gust_max or 0.0 for item in metrics), default=0.0)
        precipitation = max(
            (item.precipitation_total or 0.0 for item in metrics),
            default=0.0,
        )
        highs = [value for item in metrics if (value := item.temperature_max) is not None]
        lows = [value for item in metrics if (value := item.temperature_min) is not None]
        codes = [item.weather_code for item in metrics]
        if any(code in {95, 96, 99} for code in codes):
            highlights.append((100, f"возможна гроза {day:%d.%m}"))
        if gust >= 14:
            highlights.append((80, f"порывы до {_fmt(gust)} м/с {day:%d.%m}"))
        if precipitation >= 3:
            highlights.append((70, f"осадки до {_fmt(precipitation)} мм {day:%d.%m}"))
        if highs and max(highs) >= 30:
            highlights.append((60, f"жара до {_fmt(max(highs))} °C {day:%d.%m}"))
        if lows and min(lows) <= -15:
            highlights.append((60, f"мороз до {_fmt(min(lows))} °C {day:%d.%m}"))
    highlights.sort(key=lambda item: item[0], reverse=True)
    unique: list[str] = []
    for _, text in highlights:
        if text not in unique:
            unique.append(text)
        if len(unique) >= maximum:
            break
    return unique


def _control_points(
    points: list[ForecastPoint],
    report_dates: set[date],
) -> list[ForecastPoint]:
    if not points:
        return []
    first = points[0].valid_time_utc
    selected: list[ForecastPoint] = []
    for index, point in enumerate(points):
        if point.valid_time_local.date() not in report_dates:
            continue
        lead = (
            point.lead_hours
            if point.lead_hours is not None
            else round((point.valid_time_utc - first).total_seconds() / 3600)
        )
        interval = (
            DETAIL_INTERVAL_HOURS
            if lead <= DETAIL_SWITCH_HOUR
            else EXTENDED_DETAIL_INTERVAL_HOURS
        )
        if index == 0 or lead % interval == 0:
            selected.append(point)
    last_allowed = next(
        (
            point
            for point in reversed(points)
            if point.valid_time_local.date() in report_dates
        ),
        None,
    )
    if last_allowed is not None and last_allowed not in selected:
        selected.append(last_allowed)
    selected.sort(key=lambda point: point.valid_time_utc)
    if len(selected) <= MAX_DETAIL_ROWS:
        return selected
    stride = math.ceil(len(selected) / MAX_DETAIL_ROWS)
    reduced = selected[::stride]
    if selected[-1] not in reduced:
        reduced.append(selected[-1])
    return reduced[:MAX_DETAIL_ROWS]


def _consensus_point(
    points: list[ForecastPoint],
    reference: ForecastPoint,
) -> ForecastPoint:
    values: dict[str, ForecastValue] = {}
    for code in (
        "temperature_2m",
        "precipitation",
        "snowfall",
        "cloud_cover",
        "visibility",
        "cape",
        "wind_gusts_10m",
    ):
        sample = _values(points, code)
        if sample:
            value = (
                max(sample)
                if code
                in {
                    "precipitation",
                    "snowfall",
                    "cape",
                    "wind_gusts_10m",
                }
                else statistics.median(sample)
            )
            values[code] = ForecastValue(value=value)
    codes = [weather_presentation(point).code for point in points]
    counts = Counter(codes)
    weather_code = max(
        codes,
        key=lambda code: (counts[code], WEATHER_SEVERITY.get(code, 3)),
    )
    day_votes = sum(point.is_day is not False for point in points)
    return ForecastPoint(
        valid_time_utc=reference.valid_time_utc,
        valid_time_local=reference.valid_time_local,
        lead_hours=reference.lead_hours,
        weather_code=weather_code,
        is_day=day_votes >= len(points) / 2,
        values=values,
    )


def _detail_temperature_text(points: list[ForecastPoint]) -> str:
    values = _values(points, "temperature_2m")
    if not values:
        return "нет данных"
    centre = statistics.median(values)
    if len(values) == 1:
        return f"{_fmt(centre)} °C"
    return f"{_fmt(centre)} °C\nмодели: {_fmt(min(values))}…{_fmt(max(values))}"


def _detail_precipitation_text(points: list[ForecastPoint]) -> str:
    values = _values(points, "precipitation")
    if not values:
        return "нет данных"
    wet = sum(value >= 0.1 for value in values)
    if max(values) < 0.1:
        return "нет"
    return f"{_fmt(min(values))}–{_fmt(max(values))} мм\n{wet}/{len(values)} моделей"


def _detail_wind_text(points: list[ForecastPoint]) -> str:
    speeds = _values(points, "wind_speed_10m")
    gusts = _values(points, "wind_gusts_10m")
    directions = _values(points, "wind_direction_10m")
    if not speeds:
        return "нет данных"
    speed = statistics.median(speeds)
    direction, resultant = _circular_mean(directions)
    if direction is None or resultant < 0.2:
        direction_text = "направление различается"
    else:
        direction_text = wind_rumb(direction, speed)
    text = f"{direction_text}, {_fmt(speed)} м/с"
    if gusts:
        text += f"\nпорывы до {_fmt(max(gusts))} м/с"
    return text


def _detail_humidity_text(points: list[ForecastPoint]) -> str:
    values = _values(points, "relative_humidity_2m")
    return f"{_fmt(statistics.median(values), 0)} %" if values else "нет данных"


def _detail_pressure_text(points: list[ForecastPoint]) -> str:
    values = _values(points, "pressure_msl")
    return f"{_fmt(statistics.median(values), 0)} гПа" if values else "нет данных"


def _ensemble_temperature_text(point: ForecastPoint) -> str:
    low = _numeric(point, "temperature_2m_p10")
    high = _numeric(point, "temperature_2m_p90")
    if low is None or high is None:
        centre = _numeric(point, "temperature_2m")
        return f"{_fmt(centre)} °C" if centre is not None else "нет данных"
    return f"{_fmt(low)}…{_fmt(high)} °C"


def _ensemble_probability_text(
    points: list[ForecastPoint],
    threshold_token: str,
) -> str:
    code = f"precipitation_probability_ge_{threshold_token}mm"
    candidates = [
        point.measurement(code)
        for point in points
        if point.measurement(code) is not None
    ]
    candidates = [
        item
        for item in candidates
        if item is not None and _as_float(item.value) is not None
    ]
    if not candidates:
        return "нет данных"
    measurement = max(candidates, key=lambda item: float(item.value))
    probability = float(measurement.value)
    fraction = (
        f" ({measurement.event_count}/{measurement.sample_count})"
        if measurement.event_count is not None
        and measurement.sample_count is not None
        else ""
    )
    interval = (
        f", {measurement.accumulation_hours:g} ч"
        if measurement.accumulation_hours is not None
        else ""
    )
    return f"{_fmt(probability, 0)} %{fraction}{interval}"


def _ensemble_gust_text(points: list[ForecastPoint]) -> str:
    values = _values(points, "wind_gusts_10m_p90")
    return f"до {_fmt(max(values))} м/с" if values else "нет данных"


def _ensemble_members_text(
    points: list[ForecastPoint],
    forecast: ForecastSeries,
) -> str:
    counts = _values(points, "ensemble_member_count")
    coverage = _values(points, "ensemble_member_coverage")
    count = int(min(counts)) if counts else forecast.source.ensemble_member_count
    expected = forecast.source.ensemble_expected_member_count
    text = str(count) if count is not None else "—"
    if expected:
        text += f"/{expected}"
    if coverage:
        text += f"\n{_fmt(min(coverage), 0)} %"
    return text


def _values(points: list[ForecastPoint], code: str) -> list[float]:
    values: list[float] = []
    for point in points:
        value = _numeric(point, code)
        if value is not None and math.isfinite(value):
            values.append(value)
    return values


def _numeric(point: ForecastPoint, code: str) -> float | None:
    return _as_float(point.raw(code))


def _as_float(value) -> float | None:
    try:
        result = float(value) if value is not None else None
    except (TypeError, ValueError):
        return None
    return result if result is not None and math.isfinite(result) else None


def _circular_mean(values: list[float]) -> tuple[float | None, float]:
    if not values:
        return None, 0.0
    sine = statistics.fmean(
        math.sin(math.radians(value % 360)) for value in values
    )
    cosine = statistics.fmean(
        math.cos(math.radians(value % 360)) for value in values
    )
    resultant = math.hypot(sine, cosine)
    if resultant < 1e-6:
        return None, resultant
    return math.degrees(math.atan2(sine, cosine)) % 360, resultant


def _weekday_short(value: date) -> str:
    return ("пн", "вт", "ср", "чт", "пт", "сб", "вс")[value.weekday()]


def _fmt(value: float | None, precision: int = 1) -> str:
    if value is None:
        return "—"
    return f"{value:.{precision}f}".replace(".", ",")


def _short_model_name(forecast: ForecastSeries) -> str:
    return MODEL_SHORT_NAMES.get(forecast.source.model, forecast.source.model)
