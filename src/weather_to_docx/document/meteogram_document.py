from __future__ import annotations

import math
import statistics
import tempfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from weather_to_docx.document.compact_generator import (
    MAX_ENSEMBLE_SYSTEMS,
    MAX_REPORT_DAYS,
    MINIMUM_MODEL_COMPLETENESS,
    _configure_compact_document,
    _is_ensemble,
    _report_dates,
    _select_models,
    _short_model_name,
)
from weather_to_docx.document.compact_generator import (
    ScientificDocumentGenerator as CompactDocumentGenerator,
)
from weather_to_docx.document.styles import (
    DARK_BLUE,
    LIGHT_BLUE,
    WHITE,
    prevent_row_split,
    set_cell_shading,
    set_cell_text,
    set_cell_width,
    set_repeat_header_count,
    set_table_fixed_layout,
)
from weather_to_docx.document.weather_rules import weather_presentation
from weather_to_docx.domain.models import DocumentOptions, ForecastPoint, ForecastSeries, Location
from weather_to_docx.plotting.meteogram import MeteogramRenderer
from weather_to_docx.utils.files import safe_filename


class ScientificDocumentGenerator(CompactDocumentGenerator):
    """Компактный отчёт с профессиональными приложениями-метеограммами.

    Первые две страницы сохраняют операторскую сводку. Далее каждая пригодная
    детерминированная модель получает одну страницу с суточной таблицей и
    метеограммой на весь срок; выбранный ансамбль получает отдельную страницу
    с вероятностной таблицей и диапазонами неопределённости.
    """

    def generate(
        self,
        *,
        location: Location,
        series: list[ForecastSeries],
        options: DocumentOptions,
        output_path: Path,
    ) -> Path:
        if not options.include_meteograms:
            return super().generate(
                location=location,
                series=series,
                options=options,
                output_path=output_path,
            )
        if not series:
            raise ValueError("Для документа не передан ни один прогностический ряд")
        if any(item.location.id != location.id for item in series):
            raise ValueError("В документ нельзя объединять прогнозы для разных координат")

        deterministic = [item for item in series if not _is_ensemble(item)]
        ensembles = [item for item in series if _is_ensemble(item)]
        selection = _select_models(
            deterministic or ensembles,
            minimum_score=MINIMUM_MODEL_COMPLETENESS,
        )
        report_dates = _report_dates(selection.reference, MAX_REPORT_DAYS)
        if not report_dates:
            raise ValueError("В прогнозе нет сроков для формирования документа")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        _configure_compact_document(document)
        self._configure_properties(document, location, options)
        self._add_compact_header(
            document,
            location,
            selection,
            ensembles,
            report_dates,
            options,
        )
        self._add_highlights(document, selection, report_dates)
        self._add_daily_table(document, selection, report_dates)

        document.add_page_break()
        self._add_control_times_table(document, selection, report_dates)
        self._add_compact_notes(document, selection, ensembles)

        renderer = MeteogramRenderer(
            dpi=options.meteogram_dpi,
            smoothing=options.meteogram_smoothing,
        )
        with tempfile.TemporaryDirectory(prefix="weather-to-docx-meteograms-") as temporary:
            temporary_path = Path(temporary)
            for forecast in selection.usable:
                if _is_ensemble(forecast):
                    continue
                document.add_page_break()
                self._add_model_appendix(
                    document,
                    forecast,
                    report_dates,
                    renderer,
                    temporary_path,
                )

            if options.include_ensemble_section:
                for forecast in ensembles[:MAX_ENSEMBLE_SYSTEMS]:
                    document.add_page_break()
                    self._add_ensemble_summary(document, [forecast], report_dates)
                    image_path = temporary_path / (
                        safe_filename(f"ensemble_{forecast.source.source_id}") + ".png"
                    )
                    renderer.render_ensemble(
                        forecast,
                        image_path,
                        title=f"Ансамблевая метеограмма — {_short_model_name(forecast)}",
                    )
                    self._add_meteogram_image(
                        document,
                        image_path,
                        description=(
                            f"Ансамблевая метеограмма {_short_model_name(forecast)}: "
                            "медиана, spread и диапазон q10–q90"
                        ),
                    )
                    self._add_chart_note(
                        document,
                        "Центральная линия — медиана. Полупрозрачный внешний диапазон — "
                        "q10–q90, внутренний диапазон температуры — ±σ относительно среднего. "
                        "Вероятности осадков показаны отдельно для заданных порогов.",
                    )

            self._add_footer(document, location)
            document.save(output_path)
        return output_path

    def _add_model_appendix(
        self,
        document: Document,
        forecast: ForecastSeries,
        report_dates: list[date],
        renderer: MeteogramRenderer,
        temporary_path: Path,
    ) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run(f"Модель {_short_model_name(forecast)}")
        run.bold = True
        run.font.name = "Liberation Sans"
        run.font.size = Pt(11)

        self._add_model_daily_table(document, forecast, report_dates)
        image_path = temporary_path / (
            safe_filename(f"model_{forecast.source.source_id}") + ".png"
        )
        renderer.render_deterministic(
            forecast,
            image_path,
            title=f"Метеограмма — {_short_model_name(forecast)}",
        )
        self._add_meteogram_image(
            document,
            image_path,
            description=f"Метеограмма модели {_short_model_name(forecast)} на весь срок",
        )
        self._add_chart_note(
            document,
            "Температура, влажность, облачность, давление и ветер сглажены "
            "shape-preserving методом PCHIP. Осадки показаны нативными столбиками "
            "без сглаживания; затемнение фона соответствует ночным срокам.",
        )

    def _add_model_daily_table(
        self,
        document: Document,
        forecast: ForecastSeries,
        report_dates: list[date],
    ) -> None:
        headers = (
            "Дата",
            "Погода",
            "Температура",
            "Осадки",
            "Ветер",
            "Влажность / облачность",
            "Давление",
        )
        widths = (23, 43, 31, 31, 43, 49, 35)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(table)
        for cell, text, width in zip(table.rows[0].cells, headers, widths, strict=True):
            set_cell_text(cell, text, size=6.8, bold=True)
            set_cell_shading(cell, DARK_BLUE)
            self._set_text_color(cell, WHITE)
            set_cell_width(cell, width)
        set_repeat_header_count(table, 1)

        for row_index, day in enumerate(report_dates):
            points = [point for point in forecast.points if point.valid_time_local.date() == day]
            if not points:
                continue
            row = table.add_row()
            prevent_row_split(row)
            for cell, width in zip(row.cells, widths, strict=True):
                set_cell_width(cell, width)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_shading(cell, LIGHT_BLUE if row_index % 2 == 0 else WHITE)
            representative = _representative_point(points)
            temperature = _finite_values(points, "temperature_2m")
            precipitation = _finite_values(points, "precipitation")
            wind = _finite_values(points, "wind_speed_10m")
            gust = _finite_values(points, "wind_gusts_10m")
            humidity = _finite_values(points, "relative_humidity_2m")
            cloud = _finite_values(points, "cloud_cover")
            pressure = _finite_values(points, "pressure_msl")
            values = (
                day.strftime("%d.%m"),
                weather_presentation(representative).description,
                _range_text(temperature, "°C"),
                _sum_text(precipitation, "мм"),
                _wind_text(wind, gust),
                _humidity_cloud_text(humidity, cloud),
                _range_text(pressure, "гПа", precision=0),
            )
            for cell, text in zip(row.cells, values, strict=True):
                set_cell_text(cell, text, size=6.4)

    @staticmethod
    def _add_meteogram_image(document: Document, image_path: Path, *, description: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run()
        shape = run.add_picture(str(image_path), width=Mm(276))
        shape._inline.docPr.set("descr", description)
        shape._inline.docPr.set("title", description)

    @staticmethod
    def _add_chart_note(document: Document, text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.font.name = "Liberation Sans"
        run.font.size = Pt(6.3)
        run.font.italic = True


def _representative_point(points: list[ForecastPoint]) -> ForecastPoint:
    wettest = max(points, key=lambda point: _numeric(point, "precipitation") or 0.0)
    if (_numeric(wettest, "precipitation") or 0.0) >= 0.1:
        return wettest
    return min(points, key=lambda point: abs(point.valid_time_local.hour - 12))


def _finite_values(points: list[ForecastPoint], code: str) -> list[float]:
    values: list[float] = []
    for point in points:
        value = _numeric(point, code)
        if value is not None:
            values.append(value)
    return values


def _numeric(point: ForecastPoint, code: str) -> float | None:
    raw = point.raw(code)
    try:
        value = float(raw) if raw is not None else None
    except (TypeError, ValueError):
        return None
    return value if value is not None and math.isfinite(value) else None


def _range_text(values: list[float], unit: str, *, precision: int = 1) -> str:
    if not values:
        return "нет данных"
    low = f"{min(values):.{precision}f}".replace(".", ",")
    high = f"{max(values):.{precision}f}".replace(".", ",")
    return f"{low}…{high} {unit}"


def _sum_text(values: list[float], unit: str) -> str:
    if not values:
        return "нет данных"
    total = sum(max(0.0, value) for value in values)
    if total < 0.05:
        return "без осадков"
    return f"{total:.1f} {unit}".replace(".", ",")


def _wind_text(wind: list[float], gust: list[float]) -> str:
    if not wind:
        return "нет данных"
    text = f"до {max(wind):.1f} м/с".replace(".", ",")
    if gust:
        text += f"\nпорывы {max(gust):.1f} м/с".replace(".", ",")
    return text


def _humidity_cloud_text(humidity: list[float], cloud: list[float]) -> str:
    parts = []
    if humidity:
        parts.append(f"влажность {statistics.median(humidity):.0f} %")
    if cloud:
        parts.append(f"облачность {statistics.median(cloud):.0f} %")
    return "\n".join(parts) if parts else "нет данных"
