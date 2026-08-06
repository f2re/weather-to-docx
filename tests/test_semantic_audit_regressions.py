from __future__ import annotations

import asyncio
from datetime import UTC, date, datetime, timedelta
from pathlib import Path

from docx import Document

from weather_to_docx.analysis import build_risk_signals
from weather_to_docx.analysis.consensus import daily_precipitation_total
from weather_to_docx.analysis.impact_scales import (
    normalise_precipitation_rates,
    wind_impact_label,
)
from weather_to_docx.analysis.semantic_policy import strict_majority
from weather_to_docx.document.compact_generator import DailyModelMetrics
from weather_to_docx.document.consistent_summary import (
    build_consistent_risk_signals,
    daily_pressure_text,
)
from weather_to_docx.document.scientific_generator import (
    ScientificDocumentGenerator,
)
from weather_to_docx.document.weather_rules import derive_weather_code
from weather_to_docx.domain.models import (
    DocumentOptions,
    ForecastPoint,
    ForecastSeries,
    ForecastValue,
    Location,
    SourceMetadata,
)
from weather_to_docx.sources.demo import DemoSource
from weather_to_docx.sources.gfs_operational import GfsNomadsSource

LOCATION = Location(
    id="semantic-audit",
    name="Точка проверки",
    latitude=59.9,
    longitude=30.3,
    timezone="UTC",
)
START = datetime(2026, 8, 10, tzinfo=UTC)


def _series(
    source_id: str,
    *,
    precipitation: float = 0.0,
    weather_code: int = 2,
    include_precipitation: bool = True,
) -> ForecastSeries:
    values = {
        "temperature_2m": ForecastValue(value=20.0),
        "wind_speed_10m": ForecastValue(value=4.0),
        "wind_gusts_10m": ForecastValue(value=8.0),
        "pressure_msl": ForecastValue(value=1012.0),
        "cloud_cover": ForecastValue(value=50.0),
    }
    if include_precipitation:
        values["precipitation"] = ForecastValue(
            value=precipitation,
            source_start_step=0,
            source_end_step=3,
            accumulation_hours=3,
        )
    point = ForecastPoint(
        valid_time_utc=START + timedelta(hours=3),
        valid_time_local=START + timedelta(hours=3),
        lead_hours=3,
        weather_code=weather_code,
        is_day=False,
        values=values,
    )
    return ForecastSeries(
        location=LOCATION,
        source=SourceMetadata(
            source_id=source_id,
            provider="Test",
            model=source_id,
            product="test",
            retrieved_at_utc=START,
            exact_cycle_known=False,
        ),
        points=[point],
    )


def _text(document: Document) -> str:
    return "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )


def test_majority_is_strict_for_even_model_count() -> None:
    assert strict_majority(1) == 1
    assert strict_majority(2) == 2
    assert strict_majority(4) == 3
    assert strict_majority(6) == 4


def test_single_model_risk_has_no_inter_model_confidence() -> None:
    forecast = _series(
        "single",
        precipitation=3.0,
        weather_code=95,
    )

    signals = build_consistent_risk_signals(
        [forecast],
        [],
        [date(2026, 8, 10)],
    )

    thunder = next(
        signal for signal in signals if signal.phenomenon == "ГРОЗА"
    )
    assert thunder.scenario == "Сценарий одной модели"
    assert thunder.confidence == "не оценивается"


def test_public_analysis_api_uses_operational_risk_policy() -> None:
    forecasts = [
        _series("thunder-1", precipitation=3.0, weather_code=95),
        _series("thunder-2", precipitation=3.0, weather_code=95),
        _series("dry-1"),
        _series("dry-2"),
    ]

    signals = build_risk_signals(
        forecasts,
        [],
        [date(2026, 8, 10)],
    )

    assert all(signal.phenomenon != "ГРОЗА" for signal in signals)


def test_wind_label_does_not_turn_moderate_gust_into_strong_wind() -> None:
    assert wind_impact_label(4.0, 11.0) == "ветрено"
    assert wind_impact_label(4.0, 11.0) != "сильный ветер"
    assert wind_impact_label(10.0, 11.0) == "сильный ветер"
    assert wind_impact_label(4.0, 14.0) == "сильные порывы"


def test_weather_code_uses_interval_rate_not_raw_accumulation() -> None:
    point = ForecastPoint(
        valid_time_utc=START + timedelta(hours=6),
        valid_time_local=START + timedelta(hours=6),
        values={
            "temperature_2m": ForecastValue(value=10.0),
            "precipitation": ForecastValue(
                value=6.0,
                source_start_step=0,
                source_end_step=6,
                accumulation_hours=6,
            ),
        },
    )

    assert derive_weather_code(point) == 61


def test_gfs_cumulative_precipitation_gets_non_overlapping_intervals() -> None:
    points = []
    for end_step, accumulated in ((3, 2.0), (6, 5.0), (9, 9.0)):
        valid = START + timedelta(hours=end_step)
        points.append(
            ForecastPoint(
                valid_time_utc=valid,
                valid_time_local=valid,
                lead_hours=end_step,
                values={
                    "precipitation_accumulated": ForecastValue(
                        value=accumulated,
                        source_start_step=0,
                        source_end_step=end_step,
                    )
                },
            )
        )

    GfsNomadsSource._derive_interval_precipitation(points)

    measurements = [
        point.measurement("precipitation")
        for point in points
    ]
    assert [measurement.value for measurement in measurements] == [
        2.0,
        3.0,
        4.0,
    ]
    assert [
        measurement.source_start_step for measurement in measurements
    ] == [0, 3, 6]
    assert [
        measurement.source_end_step for measurement in measurements
    ] == [3, 6, 9]
    assert [
        measurement.accumulation_hours for measurement in measurements
    ] == [3.0, 3.0, 3.0]
    assert daily_precipitation_total(points) == 9.0

    rates = normalise_precipitation_rates(
        points,
        [point.raw("precipitation") for point in points],
    )
    assert rates == [2.0 / 3.0, 1.0, 4.0 / 3.0]


def test_pressure_main_value_is_central_not_extreme_envelope() -> None:
    source = _series("pressure")
    metrics = [
        DailyModelMetrics(
            source=source,
            weather_code=2,
            temperature_min=15.0,
            temperature_max=20.0,
            precipitation_total=0.0,
            wind_max=4.0,
            gust_max=8.0,
            pressure_min=low,
            pressure_max=high,
        )
        for low, high in (
            (990.0, 1000.0),
            (1000.0, 1010.0),
            (1005.0, 1015.0),
            (1010.0, 1020.0),
            (1015.0, 1030.0),
        )
    ]

    text = daily_pressure_text(metrics)

    assert text.splitlines()[0] == "1005–1015 гПа"
    assert "по моделям 990–1030 гПа" in text


def test_missing_precipitation_is_not_silently_treated_as_zero() -> None:
    from weather_to_docx.analysis.impact_scales import (
        daily_precipitation_summary,
    )
    from weather_to_docx.document.consistent_summary import (
        daily_precipitation_text,
    )

    dry = _series("dry", precipitation=0.0)
    missing = _series("missing", include_precipitation=False)

    assert daily_precipitation_summary(
        missing,
        date(2026, 8, 10),
    ) is None
    text = daily_precipitation_text(
        [dry, missing],
        date(2026, 8, 10),
    )
    assert "нет данных об осадках: 1/2 моделей" in text


def test_brief_document_marks_single_model_risks_as_potential(
    tmp_path: Path,
) -> None:
    forecast = asyncio.run(
        DemoSource().fetch(
            LOCATION,
            forecast_days=1,
            options={"hours": 24},
        )
    )
    output = tmp_path / "single-model.docx"

    ScientificDocumentGenerator(tmp_path / "icons").generate(
        location=LOCATION,
        series=[forecast],
        options=DocumentOptions(
            document_mode="brief",
            include_meteograms=False,
        ),
        output_path=output,
    )

    text = _text(Document(output))
    assert "Ключевые риски" in text
    assert "Сценарий одной модели" in text
    assert "уверенность: не оценивается" in text
    assert "Устойчивый сигнал" not in text
