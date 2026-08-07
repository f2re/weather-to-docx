from __future__ import annotations

import asyncio
from datetime import UTC, datetime
from pathlib import Path

from docx import Document

from weather_to_docx.document.plain_language import (
    ensemble_members_text,
    ensemble_precipitation_text,
)
from weather_to_docx.document.scientific_generator import ScientificDocumentGenerator
from weather_to_docx.domain.models import (
    DocumentOptions,
    ForecastPoint,
    ForecastSeries,
    ForecastValue,
    Location,
    SourceMetadata,
)
from weather_to_docx.sources.demo import DemoSource

LOCATION = Location(
    id="plain-language",
    name="Точка проверки",
    latitude=59.9,
    longitude=30.3,
    timezone="UTC",
)
NOW = datetime(2026, 8, 7, 6, tzinfo=UTC)


def _document_text(document: Document) -> str:
    return "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )


def test_multi_model_document_does_not_show_confidence_jargon(tmp_path: Path) -> None:
    first = asyncio.run(
        DemoSource().fetch(
            LOCATION,
            forecast_days=2,
            options={"hours": 36},
        )
    )
    second = first.model_copy(deep=True)
    second.source = second.source.model_copy(
        update={"source_id": "demo-second", "model": "Вторая модель"}
    )
    output = tmp_path / "plain.docx"

    ScientificDocumentGenerator(tmp_path / "icons").generate(
        location=LOCATION,
        series=[first, second],
        options=DocumentOptions(
            document_mode="brief",
            include_meteograms=False,
        ),
        output_path=output,
    )

    text = _document_text(Document(output))
    forbidden = (
        "Уверенность",
        "уверенность:",
        "Согласованность",
        "Устойчивый сигнал",
        "Вероятный сигнал",
        "Сценарий одной модели",
        "медиана",
    )
    for term in forbidden:
        assert term not in text
    assert "По моделям" in text
    assert "Прогноз по времени" in text


def test_single_model_document_does_not_repeat_one_of_one(tmp_path: Path) -> None:
    forecast = asyncio.run(
        DemoSource().fetch(
            LOCATION,
            forecast_days=1,
            options={"hours": 24},
        )
    )
    output = tmp_path / "single.docx"

    ScientificDocumentGenerator(tmp_path / "icons").generate(
        location=LOCATION,
        series=[forecast],
        options=DocumentOptions(
            document_mode="brief",
            include_meteograms=False,
        ),
        output_path=output,
    )

    text = _document_text(Document(output))
    assert "1 из 1 моделей" not in text
    assert "1/1 моделей" not in text
    assert "расчёт одной модели" in text
    assert "порывы: медиана" not in text
    assert "уверенность" not in text.lower()


def test_ensemble_percent_is_explained_as_share_of_variants() -> None:
    point = ForecastPoint(
        valid_time_utc=NOW,
        valid_time_local=NOW,
        values={
            "precipitation_probability_ge_0p1mm": ForecastValue(
                value=5.0,
                event_count=3,
                sample_count=51,
                accumulation_hours=3.0,
            ),
            "precipitation_probability_ge_1mm": ForecastValue(
                value=2.0,
                event_count=1,
                sample_count=51,
                accumulation_hours=3.0,
            ),
            "ensemble_member_count": ForecastValue(value=48),
        },
    )
    forecast = ForecastSeries(
        location=LOCATION,
        source=SourceMetadata(
            source_id="ensemble-test",
            provider="Test",
            model="Тестовый ансамбль",
            product="test",
            retrieved_at_utc=NOW,
            ensemble_member_count=48,
            ensemble_expected_member_count=51,
        ),
        points=[point],
    )

    text = ensemble_precipitation_text([point])

    assert "5 % вариантов (3 из 51)" in text
    assert "за 3 ч" in text
    assert "P ≥" not in text
    assert ensemble_members_text([point], forecast) == "48 из 51"


def test_web_form_hides_technical_precipitation_thresholds() -> None:
    root = Path(__file__).resolve().parents[1]
    html = (root / "src/weather_to_docx/static/index.html").read_text(encoding="utf-8")

    assert "Детерминированные модели" not in html
    assert "Пороги осадков, мм" not in html
    assert 'id="precipitationThresholds" type="hidden"' in html
    assert "Ансамбль — варианты одного прогноза" in html
    assert "Weather to DOCX 0.5.1" in html
