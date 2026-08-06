from __future__ import annotations

from pathlib import Path

import numpy as np
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm
from PIL import Image

from weather_to_docx.document.scientific_generator import ScientificDocumentGenerator
from weather_to_docx.document.verification import inspect_meteogram_docx


def _write_noisy_image(path: Path) -> None:
    random = np.random.default_rng(42)
    pixels = random.integers(0, 256, size=(420, 700, 3), dtype=np.uint8)
    Image.fromarray(pixels).save(path, format="PNG")


def _document_with_image(tmp_path: Path, *, height_mm: float) -> Path:
    image_path = tmp_path / f"chart-{height_mm:g}.png"
    _write_noisy_image(image_path)
    document = Document()
    document.add_paragraph("Метеограмма модели — проверка размеров")
    document.add_picture(
        str(image_path),
        width=Mm(276),
        height=Mm(height_mm),
    )
    output = tmp_path / f"chart-{height_mm:g}.docx"
    document.save(output)
    return output


def test_meteogram_image_with_printable_height_is_accepted(tmp_path: Path) -> None:
    inspection = inspect_meteogram_docx(
        _document_with_image(tmp_path, height_mm=136),
    )

    assert inspection.oversized_image_count == 0
    assert inspection.max_image_height_mm == 136
    assert inspection.ready


def test_meteogram_image_that_can_be_clipped_is_rejected(tmp_path: Path) -> None:
    inspection = inspect_meteogram_docx(
        _document_with_image(tmp_path, height_mm=177),
    )

    assert inspection.oversized_image_count == 1
    assert inspection.max_image_height_mm == 177
    assert not inspection.ready


def test_inline_meteogram_reserves_full_line_height(tmp_path: Path) -> None:
    image_path = tmp_path / "meteogram.png"
    _write_noisy_image(image_path)
    document = Document()

    ScientificDocumentGenerator._add_meteogram_image(
        document,
        image_path,
        description="Проверка полной высоты метеограммы",
    )

    spacing = document.paragraphs[-1]._p.pPr.spacing
    assert spacing.get(qn("w:lineRule")) == "atLeast"
    assert int(spacing.get(qn("w:line"))) >= 7900
