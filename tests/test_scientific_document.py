from __future__ import annotations

import math
import os
import re
import shutil
import subprocess
from datetime import UTC, datetime, timedelta
from pathlib import Path
from zipfile import ZipFile
from zoneinfo import ZoneInfo

import pytest
from docx import Document

from weather_to_docx.document.scientific_generator import ScientificDocumentGenerator
from weather_to_docx.domain.models import (
    DocumentOptions,
    ForecastPoint,
    ForecastSeries,
    ForecastValue,
    LeadTimeReference,
    Location,
    QualityFlag,
    SourceKind,
    SourceMetadata,
    TimezoneSource,
)

MOSCOW = ZoneInfo("Europe/Moscow")


def _location() -> Location:
    return Location(
        id="spb",
        name="Санкт-Петербург",
        latitude=59.9386,
        longitude=30.3141,
        timezone="Europe/Moscow",
        timezone_source=TimezoneSource.COORDINATES,
    )


def _deterministic(
    location: Location,
    *,
    source_id: str,
    model: str,
    offset: float = 0.0,
    incomplete: bool = False,
) -> ForecastSeries:
    start = datetime(2026, 8, 4, 0, tzinfo=UTC)
    points: list[ForecastPoint] = []
    for hour in range(168):
        valid_utc = start + timedelta(hours=hour)
        valid_local = valid_utc.astimezone(MOSCOW)
        daily_wave = 4.5 * math.sin((valid_local.hour - 6) / 24 * 2 * math.pi)
        day_index = hour // 24
        rain = 0.7 if day_index in {1, 4} and valid_local.hour in {12, 13} else 0.0
        weather_code = 61 if rain else 3 if valid_local.hour in {8, 9} else 2
        if incomplete:
            values = {
                "wind_gusts_10m": ForecastValue(value=8.0 + offset, unit="m/s"),
                "pressure_msl": ForecastValue(value=1015.0 - day_index, unit="hPa"),
                "visibility": ForecastValue(value=24_100, unit="m"),
                "cape": ForecastValue(value=30, unit="J/kg"),
                "evapotranspiration": ForecastValue(value=0.0, unit="mm"),
            }
        else:
            speed = 3.0 + 0.7 * math.sin(hour / 10) + offset * 0.1
            values = {
                "temperature_2m": ForecastValue(
                    value=17.0 + daily_wave + day_index * 0.4 + offset,
                    unit="°C",
                ),
                "apparent_temperature": ForecastValue(
                    value=16.5 + daily_wave + day_index * 0.4 + offset,
                    unit="°C",
                ),
                "dew_point_2m": ForecastValue(value=11.0 + offset * 0.2, unit="°C"),
                "relative_humidity_2m": ForecastValue(
                    value=70 - daily_wave * 2,
                    unit="%",
                ),
                "precipitation": ForecastValue(
                    value=rain,
                    unit="mm",
                    accumulation_hours=1,
                ),
                "rain": ForecastValue(value=rain, unit="mm"),
                "wind_speed_10m": ForecastValue(value=speed, unit="m/s"),
                "wind_direction_10m": ForecastValue(
                    value=(220 + offset * 5 + hour) % 360,
                    unit="°",
                ),
                "wind_gusts_10m": ForecastValue(value=speed + 4.0, unit="m/s"),
                "pressure_msl": ForecastValue(
                    value=1018.0 - day_index * 1.2 + offset * 0.2,
                    unit="hPa",
                ),
                "surface_pressure": ForecastValue(value=1017.0, unit="hPa"),
                "cloud_cover": ForecastValue(value=80 if rain else 45, unit="%"),
                "visibility": ForecastValue(value=18_000, unit="m"),
                "cape": ForecastValue(value=250 if rain else 20, unit="J/kg"),
                "shortwave_radiation": ForecastValue(value=450, unit="W/m²"),
                "soil_temperature_0cm": ForecastValue(value=18.0, unit="°C"),
                "et0_fao_evapotranspiration": ForecastValue(value=0.2, unit="mm"),
            }
        points.append(
            ForecastPoint(
                valid_time_utc=valid_utc,
                valid_time_local=valid_local,
                lead_hours=hour,
                weather_code=weather_code,
                is_day=6 <= valid_local.hour <= 21,
                values=values,
            )
        )
    return ForecastSeries(
        location=location,
        source=SourceMetadata(
            source_id=source_id,
            provider="Open-Meteo / Test",
            model=model,
            product="hourly point forecast",
            source_kind=SourceKind.DETERMINISTIC,
            cycle_time_utc=None,
            retrieved_at_utc=start,
            horizon_hours=167,
            native_time_step_hours=1,
            lead_time_reference=LeadTimeReference.RESPONSE_START,
            exact_cycle_known=False,
            attribution="English attribution that must not enter the operator report",
        ),
        points=points,
        warnings=["Technical source warning that must not be repeated."],
    )


def _ensemble(location: Location) -> ForecastSeries:
    start = datetime(2026, 8, 4, 0, tzinfo=UTC)
    calculated = QualityFlag.CALCULATED
    points: list[ForecastPoint] = []
    for hour in range(0, 168, 6):
        valid_utc = start + timedelta(hours=hour)
        valid_local = valid_utc.astimezone(MOSCOW)
        day_index = hour // 24
        probability = 65 if day_index in {1, 4} else 10
        values = {
            "temperature_2m": ForecastValue(value=18.0 + day_index * 0.4, unit="°C"),
            "temperature_2m_p10": ForecastValue(
                value=15.0 + day_index * 0.4,
                unit="°C",
                quality=calculated,
                sample_count=31,
            ),
            "temperature_2m_p90": ForecastValue(
                value=22.0 + day_index * 0.4,
                unit="°C",
                quality=calculated,
                sample_count=31,
            ),
            "precipitation_probability_ge_0p1mm": ForecastValue(
                value=probability,
                unit="%",
                quality=calculated,
                event_count=20 if probability >= 60 else 3,
                sample_count=31,
                accumulation_hours=6,
            ),
            "precipitation_probability_ge_1mm": ForecastValue(
                value=35 if probability >= 60 else 0,
                unit="%",
                quality=calculated,
                event_count=11 if probability >= 60 else 0,
                sample_count=31,
                accumulation_hours=6,
            ),
            "wind_gusts_10m_p90": ForecastValue(
                value=12.0 + day_index * 0.3,
                unit="m/s",
                quality=calculated,
                sample_count=31,
            ),
            "ensemble_member_count": ForecastValue(
                value=30,
                quality=QualityFlag.SUSPECT,
                sample_count=30,
            ),
            "ensemble_member_coverage": ForecastValue(
                value=97,
                unit="%",
                quality=QualityFlag.SUSPECT,
                sample_count=30,
            ),
        }
        points.append(
            ForecastPoint(
                valid_time_utc=valid_utc,
                valid_time_local=valid_local,
                lead_hours=hour,
                weather_code=61 if probability >= 60 else 2,
                is_day=True,
                values=values,
            )
        )
    return ForecastSeries(
        location=location,
        source=SourceMetadata(
            source_id="open_meteo_gefs_0p25",
            provider="Open-Meteo / NOAA",
            model="NOAA GEFS 0.25°",
            product="ensemble distribution",
            source_kind=SourceKind.ENSEMBLE,
            retrieved_at_utc=start,
            exact_cycle_known=False,
            lead_time_reference=LeadTimeReference.RESPONSE_START,
            ensemble_member_count=30,
            ensemble_expected_member_count=31,
            ensemble_member_coverage_percent=97,
            member_weighting="equal",
            probability_calibration="raw_uncalibrated_member_fraction",
        ),
        points=points,
    )


def _report_series(location: Location) -> list[ForecastSeries]:
    return [
        _deterministic(
            location,
            source_id="open_meteo_gfs",
            model="NOAA GFS 0.25°",
            incomplete=True,
        ),
        _deterministic(
            location,
            source_id="open_meteo_ecmwf_ifs",
            model="ECMWF IFS 0.25° Open Data",
            offset=0.0,
        ),
        _deterministic(
            location,
            source_id="open_meteo_dwd_icon_global",
            model="DWD ICON Global",
            offset=1.1,
        ),
        _deterministic(
            location,
            source_id="open_meteo_gem_gdps",
            model="ECCC GEM Global (GDPS)",
            offset=-0.8,
        ),
        _ensemble(location),
    ]


def _generate_report(tmp_path: Path) -> Path:
    location = _location()
    output = tmp_path / "compact-report.docx"
    ScientificDocumentGenerator(tmp_path / "icons").generate(
        location=location,
        series=_report_series(location),
        options=DocumentOptions(
            page_size="A4",
            parameter_profile="operational",
            include_all_parameters=False,
            include_meteograms=True,
            meteogram_dpi=120,
        ),
        output_path=output,
    )
    return output


def _all_text(document: Document) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*paragraphs, *cells])


def test_report_contains_summary_and_model_meteogram_appendices(tmp_path: Path) -> None:
    output = _generate_report(tmp_path)
    document = Document(output)
    text = _all_text(document)

    assert len(document.tables) == 6
    assert "Прогноз по дням" in text
    assert "Прогноз по контрольным срокам" in text
    assert "Согласованность" in text
    assert "Не использованы в сводке из-за неполных данных: NOAA GFS" in text
    assert "Модель ECMWF IFS" in text
    assert "Модель ICON" in text
    assert "Модель GDPS" in text
    assert "Неопределённость прогноза" in text

    assert "1. Наглядный прогноз" not in text
    assert "2. Подробный метеорологический отчёт" not in text
    assert "hourly point forecast" not in text
    assert "Атрибуция" not in text
    assert "Условия использования" not in text
    assert "MSL" not in text
    assert "SFC" not in text
    assert "CAPE" not in text
    assert "CIN" not in text
    assert "VPD" not in text
    assert "ET₀" not in text
    assert "Почва" not in text
    assert "Радиация" not in text

    assert "m/s" not in text
    assert re.search(r"(?<![А-Яа-яЁё])mm(?![А-Яа-яЁё])", text) is None
    assert "hPa" not in text
    assert "м/с" in text
    assert "мм" in text
    assert "гПа" in text

    inline_shapes = document.inline_shapes
    assert len(inline_shapes) == 4
    assert all(shape.width.mm > 250 for shape in inline_shapes)
    assert all(shape.height.mm > 80 for shape in inline_shapes)

    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert "Метеограмма модели ECMWF IFS" in xml
    assert "Метеограмма модели ICON" in xml
    assert "Метеограмма модели GDPS" in xml
    assert "Ансамблевая метеограмма GEFS" in xml
    assert len(media) >= 4


def test_report_without_meteograms_remains_two_pages_structure(tmp_path: Path) -> None:
    location = _location()
    output = tmp_path / "without-meteograms.docx"
    ScientificDocumentGenerator(tmp_path / "icons").generate(
        location=location,
        series=_report_series(location),
        options=DocumentOptions(
            include_meteograms=False,
            page_size="A4",
            parameter_profile="operational",
        ),
        output_path=output,
    )
    document = Document(output)
    assert len(document.tables) == 3
    assert len(document.inline_shapes) == 0


def test_incomplete_only_source_is_rejected_instead_of_printing_blanks(
    tmp_path: Path,
) -> None:
    location = _location()
    with pytest.raises(ValueError, match="не содержит достаточного набора ключевых полей"):
        ScientificDocumentGenerator(tmp_path / "icons").generate(
            location=location,
            series=[
                _deterministic(
                    location,
                    source_id="open_meteo_gfs",
                    model="NOAA GFS 0.25°",
                    incomplete=True,
                )
            ],
            options=DocumentOptions(),
            output_path=tmp_path / "invalid.docx",
        )


@pytest.mark.skipif(
    not shutil.which("libreoffice") or not shutil.which("pdfinfo"),
    reason="LibreOffice и pdfinfo нужны для проверки физической пагинации",
)
def test_rendered_report_page_count_matches_appendices(tmp_path: Path) -> None:
    output = _generate_report(tmp_path)
    profile = tmp_path / "libreoffice-profile"
    profile.mkdir()
    environment = os.environ.copy()
    environment["HOME"] = str(tmp_path)
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            str(output),
        ],
        check=True,
        capture_output=True,
        text=True,
        timeout=120,
        env=environment,
    )
    pdf = output.with_suffix(".pdf")
    info = subprocess.run(
        ["pdfinfo", str(pdf)],
        check=True,
        capture_output=True,
        text=True,
        timeout=30,
    ).stdout
    match = re.search(r"^Pages:\s+(\d+)$", info, re.MULTILINE)
    assert match is not None
    assert int(match.group(1)) == 6
