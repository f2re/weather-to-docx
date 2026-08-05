from __future__ import annotations

from datetime import UTC, datetime, timedelta
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document

from weather_to_docx.document.scientific_generator import ScientificDocumentGenerator
from weather_to_docx.domain.models import (
    DocumentOptions,
    ForecastPoint,
    ForecastSeries,
    ForecastValue,
    Location,
    SourceMetadata,
)

MOSCOW = ZoneInfo("Europe/Moscow")


def test_visible_document_labels_are_russian(tmp_path: Path) -> None:
    location = Location(
        id="spb-russian",
        name="Санкт-Петербург",
        latitude=59.9607,
        longitude=30.1587,
        timezone="Europe/Moscow",
    )
    start = datetime(2026, 8, 5, tzinfo=UTC)
    points = []
    for hour in range(0, 48, 3):
        valid_utc = start + timedelta(hours=hour)
        local = valid_utc.astimezone(MOSCOW)
        points.append(
            ForecastPoint(
                valid_time_utc=valid_utc,
                valid_time_local=local,
                lead_hours=hour,
                weather_code=2,
                is_day=None,
                values={
                    "temperature_2m": ForecastValue(value=18 + hour / 24),
                    "relative_humidity_2m": ForecastValue(value=70),
                    "precipitation": ForecastValue(value=0.0),
                    "wind_speed_10m": ForecastValue(value=3.0),
                    "wind_direction_10m": ForecastValue(value=240.0),
                    "wind_gusts_10m": ForecastValue(value=6.0),
                    "pressure_msl": ForecastValue(value=1015.0),
                    "cloud_cover": ForecastValue(value=50.0),
                },
            )
        )
    forecast = ForecastSeries(
        location=location,
        source=SourceMetadata(
            source_id="gfs-russian-test",
            provider="NOAA",
            model="Global Forecast System (GFS)",
            product="forecast",
            retrieved_at_utc=start,
            exact_cycle_known=False,
        ),
        points=points,
    )

    output = tmp_path / "russian.docx"
    ScientificDocumentGenerator(tmp_path / "icons").generate(
        location=location,
        series=[forecast],
        options=DocumentOptions(include_meteograms=True),
        output_path=output,
    )

    document = Document(output)
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    assert "NOAA GFS" in text
    assert "Global Forecast System" not in text
    assert "Europe/Moscow" not in text
    assert "UTC+03:00" in text
    assert "shape-preserving" not in text
    assert "формосохраняющим методом PCHIP" not in text
