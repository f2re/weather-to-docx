from __future__ import annotations

from datetime import UTC, datetime, timedelta
from pathlib import Path

from docx import Document

from weather_to_docx.document.scientific_generator import ScientificDocumentGenerator
from weather_to_docx.domain.models import (
    DocumentOptions,
    ForecastPoint,
    ForecastSeries,
    ForecastValue,
    Location,
    SourceMetadata,
)

LOCATION = Location(
    id="consistent-document",
    name="Санкт-Петербург",
    latitude=59.9386,
    longitude=30.3141,
    timezone="UTC",
)
START = datetime(2026, 8, 11, tzinfo=UTC)


def _series(source_id: str, total_mm: float, weather_code: int) -> ForecastSeries:
    points = []
    for index, hour in enumerate((6, 12, 18, 23)):
        valid = START + timedelta(hours=hour)
        points.append(
            ForecastPoint(
                valid_time_utc=valid,
                valid_time_local=valid,
                lead_hours=hour,
                weather_code=weather_code,
                is_day=6 <= hour <= 18,
                values={
                    "temperature_2m": ForecastValue(value=16.0 + index),
                    "relative_humidity_2m": ForecastValue(value=78.0),
                    "precipitation": ForecastValue(
                        value=total_mm / 4,
                        source_start_step=index * 6,
                        source_end_step=(index + 1) * 6,
                        accumulation_hours=6,
                    ),
                    "wind_speed_10m": ForecastValue(value=5.0),
                    "wind_direction_10m": ForecastValue(value=270.0),
                    "wind_gusts_10m": ForecastValue(value=10.0),
                    "pressure_msl": ForecastValue(value=1008.0),
                    "cloud_cover": ForecastValue(value=80.0),
                },
            )
        )
    return ForecastSeries(
        location=LOCATION,
        source=SourceMetadata(
            source_id=source_id,
            provider="Test",
            model=source_id,
            product="test",
            retrieved_at_utc=START,
            exact_cycle_known=False,
        ),
        points=points,
    )


def _document_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )


def test_generated_document_uses_central_value_not_worst_model(tmp_path: Path) -> None:
    forecasts = [
        _series("GFS", 2.5, 61),
        _series("ICON", 3.9, 61),
        _series("AIFS", 11.2, 61),
        _series("IFS", 15.0, 61),
        _series("GDPS", 38.2, 63),
    ]
    output = tmp_path / "consistent.docx"

    ScientificDocumentGenerator(tmp_path / "icons").generate(
        location=LOCATION,
        series=forecasts,
        options=DocumentOptions(
            document_mode="brief",
            include_meteograms=False,
        ),
        output_path=output,
    )
    text = _document_text(output)

    assert "обычно около 11,2 мм за сутки" in text
    assert "по моделям 2,5–38,2 мм" in text
    assert "медиана" not in text
    assert "СИЛЬНЫЕ ОСАДКИ" not in text
    assert "Сильный дождь" not in text
    assert "очень много осадков" not in text
