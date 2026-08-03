from __future__ import annotations

import asyncio
import zipfile
from pathlib import Path

from docx import Document

from weather_to_docx.document.generator import DocumentGenerator
from weather_to_docx.domain.models import DocumentOptions, Location
from weather_to_docx.sources.demo import DemoSource


def test_document_contains_two_tables_per_source(tmp_path: Path) -> None:
    location = Location(
        id="test",
        name="Тестовая точка",
        latitude=59.94,
        longitude=30.31,
        elevation_m=10,
        timezone="Europe/Moscow",
    )
    series = asyncio.run(DemoSource().fetch(location, 1, {"hours": 12}))
    output = tmp_path / "forecast.docx"
    DocumentGenerator(tmp_path / "icons").generate(
        location=location,
        series=[series],
        options=DocumentOptions(),
        output_path=output,
    )

    assert output.is_file()
    with zipfile.ZipFile(output) as archive:
        assert "word/document.xml" in archive.namelist()
        assert any(name.startswith("word/media/") for name in archive.namelist())
        document_xml = archive.read("word/document.xml").decode("utf-8")
        assert "Пиктограмма погоды:" in document_xml
    document = Document(output)
    assert len(document.tables) == 4  # title metadata, source metadata, summary, detailed
    assert "Наглядный прогноз" in "\n".join(paragraph.text for paragraph in document.paragraphs)
