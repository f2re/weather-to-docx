from __future__ import annotations

import io
import json
import zipfile
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from weather_to_docx.api.app import create_app
from weather_to_docx.services.worker import run_worker
from weather_to_docx.settings import Settings


def _location_payload(
    identifier: str,
    name: str,
    latitude: float,
    longitude: float,
) -> dict:
    return {
        "id": identifier,
        "name": name,
        "latitude": latitude,
        "longitude": longitude,
        "elevation_m": 12,
        "timezone": "Europe/Moscow",
        "group": "Приёмочные точки",
    }


def test_persistent_operator_workflow_after_restarts(tmp_path: Path) -> None:
    settings = Settings(data_dir=tmp_path / "data")
    locations = [
        _location_payload("point-one", "Тестовая точка 1", 59.9386, 30.3141),
        _location_payload("point-two", "Тестовая точка 2", 55.7558, 37.6176),
    ]

    with TestClient(create_app(settings)) as client:
        for location in locations:
            response = client.post("/api/v1/locations", json=location)
            assert response.status_code == 201
        export_response = client.get("/api/v1/locations/export")
        assert export_response.status_code == 200
        assert {item["id"] for item in export_response.json()} == {
            "point-one",
            "point-two",
        }

    with TestClient(create_app(settings)) as client:
        stored_locations = client.get("/api/v1/locations?limit=100").json()
        assert [item["id"] for item in stored_locations] == [
            "point-one",
            "point-two",
        ]
        response = client.post(
            "/api/v1/jobs",
            json={
                "batch_name": "operator-acceptance",
                "locations": stored_locations,
                "sources": [
                    {
                        "source_id": "demo",
                        "forecast_days": 1,
                        "options": {"hours": 12},
                    }
                ],
                "document": {
                    "title": "Приёмочный метеорологический прогноз",
                    "page_size": "A4",
                    "include_detailed_table": True,
                    "include_all_parameters": False,
                    "parameter_profile": "operational",
                },
            },
        )
        assert response.status_code == 201
        job_id = response.json()["id"]

    assert run_worker(settings, once=True) == 1

    with TestClient(create_app(settings)) as client:
        response = client.get(f"/api/v1/jobs/{job_id}")
        assert response.status_code == 200
        job = response.json()
        assert job["status"] == "completed"
        artifacts = job["result"]["artifacts"]
        assert [artifact["kind"] for artifact in artifacts].count("docx") == 2
        assert {artifact["kind"] for artifact in artifacts} == {
            "docx",
            "manifest",
            "zip",
        }

        docx_artifacts = [
            (index, artifact)
            for index, artifact in enumerate(artifacts)
            if artifact["kind"] == "docx"
        ]
        for artifact_index, artifact in docx_artifacts:
            download = client.get(
                f"/api/v1/jobs/{job_id}/artifacts/{artifact_index}"
            )
            assert download.status_code == 200
            assert download.content.startswith(b"PK")
            document = Document(io.BytesIO(download.content))
            # Без ансамбля компактный документ содержит суточную и срочную таблицы.
            assert len(document.tables) == 2
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            expected_name = next(
                location["name"]
                for location in locations
                if location["id"] == artifact["location_id"]
            )
            assert expected_name in text
            assert "Прогноз по дням" in text
            assert "Прогноз по контрольным срокам" in text
            assert "Подробный метеорологический отчёт" not in text

        manifest_index = next(
            index
            for index, artifact in enumerate(artifacts)
            if artifact["kind"] == "manifest"
        )
        manifest_response = client.get(
            f"/api/v1/jobs/{job_id}/artifacts/{manifest_index}"
        )
        manifest = json.loads(manifest_response.content)
        assert manifest["batch_id"] == job_id
        assert len(manifest["locations"]) == 2
        assert len(manifest["artifacts"]) == 2

        zip_index = next(
            index
            for index, artifact in enumerate(artifacts)
            if artifact["kind"] == "zip"
        )
        zip_response = client.get(
            f"/api/v1/jobs/{job_id}/artifacts/{zip_index}"
        )
        with zipfile.ZipFile(io.BytesIO(zip_response.content)) as archive:
            names = archive.namelist()
            assert "manifest.json" in names
            assert len([name for name in names if name.endswith(".docx")]) == 2

        retry_response = client.post(f"/api/v1/jobs/{job_id}/retry")
        assert retry_response.status_code == 201
        retry_id = retry_response.json()["id"]
        cancel_response = client.post(f"/api/v1/jobs/{retry_id}/cancel")
        assert cancel_response.status_code == 200
        assert cancel_response.json()["status"] == "cancelled"
