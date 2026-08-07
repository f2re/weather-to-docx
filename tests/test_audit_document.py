from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document

from weather_to_docx.document.scientific_generator import ScientificDocumentGenerator
from weather_to_docx.document.verification import inspect_meteogram_docx
from weather_to_docx.domain.models import DocumentOptions, Location
from weather_to_docx.sources.demo import DemoSource


LOCATION = Location(
    id="audit-document",
    name="Санкт-Петербург",
    latitude=59.9386,
    longitude=30.3141,
    timezone="Europe/Moscow",
)


def _forecast(hours: int = 48):
    return asyncio.run(
        DemoSource().fetch(
            LOCATION,
            forecast_days=max(1, (hours + 23) // 24),
            options={"hours": hours},
        )
    )


def _text(document: Document) -> str:
    return "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )


def _large_meteograms(document: Document):
    return [
        shape
        for shape in document.inline_shapes
        if shape.width.mm > 250 and shape.height.mm > 80
    ]


def test_brief_single_model_hides_false_agreement_column(tmp_path: Path) -> None:
    output = tmp_path / "brief.docx"
    ScientificDocumentGenerator(tmp_path / "icons").generate(
        location=LOCATION,
        series=[_forecast(48)],
        options=DocumentOptions(
            document_mode="brief",
            include_meteograms=False,
        ),
        output_path=output,
    )
    document = Document(output)
    text = _text(document)
    assert "Ключевые риски" in text
    assert "Согласованность" not in text
    assert "одна модель" not in text
    assert "Прогноз по времени" in text
    assert not _large_meteograms(document)


def test_short_expert_document_uses_dynamic_layout_and_graph_only_page(
    tmp_path: Path,
) -> None:
    output = tmp_path / "expert.docx"
    ScientificDocumentGenerator(tmp_path / "icons").generate(
        location=LOCATION,
        series=[_forecast(48)],
        options=DocumentOptions(
            document_mode="expert",
            include_meteograms=True,
            meteogram_dpi=120,
        ),
        output_path=output,
    )
    inspection = inspect_meteogram_docx(output)
    document = Document(output)
    text = _text(document)
    assert inspection.structured_page_count == 2
    assert inspection.large_media_count >= 1
    assert inspection.has_risk_section
    assert "Метеограмма модели Синтетические данные" in text
    assert "формосохраняющим методом PCHIP" not in text
    assert "Расчёт модели:" in text
    assert "Цикл:" not in text
    assert len(document.tables) == 3


def test_full_mode_keeps_additional_model_table_on_separate_page(tmp_path: Path) -> None:
    output = tmp_path / "full.docx"
    ScientificDocumentGenerator(tmp_path / "icons").generate(
        location=LOCATION,
        series=[_forecast(48)],
        options=DocumentOptions(
            document_mode="full",
            page_size="A3",
            parameter_profile="extended",
            include_meteograms=True,
            meteogram_dpi=120,
        ),
        output_path=output,
    )
    inspection = inspect_meteogram_docx(output)
    document = Document(output)
    assert inspection.structured_page_count == 3
    assert len(document.tables) == 4


def test_document_option_defaults_to_expert_mode() -> None:
    options = DocumentOptions()
    assert options.document_mode == "expert"
    assert options.include_meteograms is True
