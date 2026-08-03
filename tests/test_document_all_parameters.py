from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

from docx import Document

from weather_to_docx.document.generator import DocumentGenerator
from weather_to_docx.domain.models import (
    DocumentOptions,
    ForecastPoint,
    ForecastSeries,
    ForecastValue,
    Location,
    QualityFlag,
    SourceMetadata,
)


def test_second_table_contains_ensemble_and_vertical_parameters(tmp_path: Path) -> None:
    location = Location(
        id="ensemble-point",
        name="Ансамблевая точка",
        latitude=59.94,
        longitude=30.31,
        timezone="Europe/Moscow",
    )
    valid_utc = datetime(2026, 8, 3, 0, tzinfo=UTC)
    point = ForecastPoint(
        valid_time_utc=valid_utc,
        valid_time_local=valid_utc,
        lead_hours=0,
        weather_code=2,
        is_day=False,
        values={
            "temperature_2m": ForecastValue(value=18.0, unit="°C"),
            "temperature_2m_spread": ForecastValue(
                value=1.8,
                unit="°C",
                quality=QualityFlag.CALCULATED,
            ),
            "temperature_2m_p10": ForecastValue(
                value=15.9,
                unit="°C",
                quality=QualityFlag.CALCULATED,
            ),
            "temperature_2m_p90": ForecastValue(
                value=20.4,
                unit="°C",
                quality=QualityFlag.CALCULATED,
            ),
            "ensemble_member_count": ForecastValue(
                value=31,
                quality=QualityFlag.CALCULATED,
            ),
            "temperature_850hPa": ForecastValue(value=8.4, unit="°C"),
            "wind_speed_500hPa": ForecastValue(value=22.5, unit="м/с"),
            "custom_provider_field": ForecastValue(value=42, unit="ед."),
        },
    )
    series = ForecastSeries(
        location=location,
        source=SourceMetadata(
            source_id="test_ensemble",
            provider="Тестовый поставщик",
            model="Тестовый ансамбль",
            product="ensemble statistics",
            retrieved_at_utc=valid_utc,
            horizon_hours=0,
            native_time_step_hours=3,
            exact_cycle_known=False,
            ensemble_member_count=31,
            upstream_model_id="test_ensemble_31",
        ),
        points=[point],
    )

    output = tmp_path / "all-parameters.docx"
    DocumentGenerator(tmp_path / "icons").generate(
        location=location,
        series=[series],
        options=DocumentOptions(include_all_parameters=True),
        output_path=output,
    )

    document = Document(output)
    assert len(document.tables) == 4
    detailed_text = "\n".join(
        cell.text
        for row in document.tables[-1].rows
        for cell in row.cells
    )
    assert "Все дополнительные параметры" in detailed_text
    assert "T σ" in detailed_text
    assert "Члены N" in detailed_text
    assert "Температура на 850 гПа" in detailed_text
    assert "Скорость ветра на 500 гПа" in detailed_text
    assert "custom_provider_field" in detailed_text
    assert "1.80* °C" in detailed_text


def test_additional_column_can_be_disabled(tmp_path: Path) -> None:
    location = Location(
        id="simple-point",
        name="Простая точка",
        latitude=55.75,
        longitude=37.62,
        timezone="Europe/Moscow",
    )
    valid_utc = datetime(2026, 8, 3, 0, tzinfo=UTC)
    series = ForecastSeries(
        location=location,
        source=SourceMetadata(
            source_id="test",
            provider="Тест",
            model="Тест",
            product="test",
            retrieved_at_utc=valid_utc,
        ),
        points=[
            ForecastPoint(
                valid_time_utc=valid_utc,
                valid_time_local=valid_utc,
                values={
                    "temperature_2m": ForecastValue(value=20, unit="°C"),
                    "temperature_2m_spread": ForecastValue(value=1, unit="°C"),
                },
            )
        ],
    )

    output = tmp_path / "operational.docx"
    DocumentGenerator(tmp_path / "icons").generate(
        location=location,
        series=[series],
        options=DocumentOptions(include_all_parameters=False),
        output_path=output,
    )
    document = Document(output)
    detailed_text = "\n".join(
        cell.text
        for row in document.tables[-1].rows
        for cell in row.cells
    )
    assert "Все дополнительные параметры" not in detailed_text
    assert "T σ" not in detailed_text
